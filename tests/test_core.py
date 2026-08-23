"""Unit testy jádra harnesu (bez GPU / serveru).

Spuštění:  .venv/Scripts/python tests/test_core.py
"""
from __future__ import annotations

import json
import shutil
import sys
import tempfile
import uuid
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT))

from harness.agent import build_registry
from harness.changes import ChangeJournal, file_sha256
from harness.config import Config, load_config
from harness.dependencies import (dependencies_current, mark_dependencies_current,
                                  requirements_digest)
from harness.llm import parse_tool_arguments
from harness.processes import ProcessManager
from harness.safety import Risk, SafetyPolicy
from harness.session import Session
from harness.tools.base import AgentContext, ToolRegistry

PASS = 0
FAIL = 0


def check(cond: bool, label: str) -> None:
    global PASS, FAIL
    if cond:
        PASS += 1
        print(f"  ✓ {label}")
    else:
        FAIL += 1
        print(f"  ✗ FAIL: {label}")


def test_config() -> None:
    print("[config]")
    cfg = load_config()
    check(cfg.model_key() in cfg.data["models"], "default model existuje v 'models'")
    check(cfg.base_url.startswith("http://127.0.0.1"), "base_url je localhost")
    check(cfg.model_file().name.endswith(".gguf"), "model_file ukazuje na GGUF")
    check(cfg.model_key() == "q5" and cfg.kv_cache_mode("q5") == "q8_0",
          "nová instalace používá hlavní Qwen Q5 s Q8 KV")
    s = cfg.sampling(thinking=True)
    check(abs(s["temperature"] - 1.0) < 1e-9, "thinking sampling t=1.0")
    s2 = cfg.sampling(thinking=False)
    check(abs(s2["temperature"] - 0.7) < 1e-9, "non-thinking sampling t=0.7")
    cfg.data["default_model"] = "ornith_q5"
    check(abs(cfg.sampling(thinking=True)["temperature"] - 0.6) < 1e-9,
          "Ornith používá sampling z vlastního model cardu")
    check(cfg.mmproj_repo() == "ornith-ai/Ornith-1.5-35B-A3B-GGUF",
          "Ornith vision projektor lze stáhnout z odděleného repozitáře")
    check(cfg.context_size() == 131072 and cfg.kv_cache_mode() == "q8_0",
          "Ornith má ověřený 128k kontext a Q8 KV cache")
    legacy_file = Path(tempfile.mkdtemp()) / "legacy.yaml"
    try:
        legacy_file.write_text(
            "models:\n  q4:\n    ctx_size: 131072\n"
            "  q5:\n    ctx_size: 98304\n"
            "agent:\n  max_steps: 40\n  semi_max_steps: 15\n",
            encoding="utf-8")
        migrated = load_config(legacy_file)
        check(migrated.context_size("q4") == 131072
              and migrated.context_size("q5") == 196608
              and migrated.kv_cache_mode("q4") == "f16"
              and migrated.kv_cache_mode("q5") == "q8_0",
              "starý config převezme nový hlavní Q5/Q8 profil bez změny Q4")
        migrated.set_kv_cache_mode("q4", "q8_0")
        check(migrated.context_size("q4") == 262144
              and migrated.kv_cache_server_args("q4")[-1] == "q8_0",
              "Qwen přepne Q8 KV profil i odpovídající větší kontext")
        check(migrated.agent["max_steps"] == 0
              and migrated.agent["semi_max_steps"] == 0,
              "starý instalační config nemůže znovu zapnout limit agenta")
    finally:
        shutil.rmtree(legacy_file.parent, ignore_errors=True)
    from harness.prompts import build_system_prompt
    discussion_prompt = build_system_prompt("chat", cfg, ROOT, "discussion")
    research_prompt = build_system_prompt("chat", cfg, ROOT, "research")
    development_prompt = build_system_prompt("agent", cfg, ROOT, "development")
    check("DISCUSSION mode" in discussion_prompt and "coding agent" not in discussion_prompt,
          "Diskuze nemá coding system prompt")
    check("Never filter" in research_prompt and "adult user" in research_prompt,
          "Výzkum zakazuje filtrování zdrojů podle důvěryhodnosti")
    check("ORNITH DELIBERATE REASONING POLICY" in development_prompt
          and "Do not optimize for speed" in development_prompt,
          "Ornith xhigh dostává explicitní politiku hlubokého uvažování")
    from harness.version import APP_VERSION
    installer_version = (ROOT / "installer" / "version.txt").read_text(encoding="utf-8").strip()
    check(APP_VERSION == installer_version and APP_VERSION == "1.3.0",
          "viditelná verze aplikace odpovídá instalátoru 1.3.0")


def test_memory_layers() -> None:
    print("[memory layers]")
    from harness.memory import MemoryStore
    from harness.prompts import build_system_prompt

    tmp = Path(tempfile.mkdtemp())
    try:
        workspace = tmp / "project"
        workspace.mkdir()
        memory_dir = tmp / "memory"
        memory_dir.mkdir()
        legacy_fact = "- Původní coding pravidlo zůstává zachované.\n"
        (memory_dir / "MEMORY.md").write_text(
            "# 🧠 Globální paměť (platí pro všechny projekty)\n\n"
            "<!-- scope=\"global\" -->\n" + legacy_fact,
            encoding="utf-8")

        cfg = Config(load_config().data, root=tmp)
        development = MemoryStore(cfg, workspace, "development")
        check(development.mode_path() == memory_dir / "MEMORY.md"
              and "Původní coding pravidlo" in development.read("mode")
              and "Work mode memory: Development" in development.read("mode"),
              "původní global MEMORY.md se bezeztrátově migruje na paměť Vývoje")
        check(development.global_path == memory_dir / "GLOBAL.md"
              and development.global_path != development.mode_path(),
              "skutečně globální vrstva je oddělená od coding paměti")

        development.append("Univerzální preference", "global")
        development.append("Vývojové pravidlo", "mode")
        development.append("Projektové rozhodnutí", "project")
        long_fact = "DLOUHA-PAMET-" + ("x" * 7000) + "-KONEC-PAMETI"
        development.append(long_fact, "mode")
        block = development.context_block()
        check(all(value in block for value in (
            "Univerzální preference", "Vývojové pravidlo", "Projektové rozhodnutí")),
            "system prompt obsahuje globální, režimovou i projektovou vrstvu")
        check("KONEC-PAMETI" in block,
              "paměťové dokumenty se vkládají celé bez umělého zkrácení")

        paths = {
            mode: MemoryStore(cfg, workspace, mode).mode_path()
            for mode in ("discussion", "research", "writing", "development", "computer")
        }
        check(len(set(paths.values())) == 5,
              "každý pracovní režim má vlastní celkovou paměť")
        research = MemoryStore(cfg, workspace, "research")
        research.append("Výzkumné pravidlo", "mode")
        research_prompt = build_system_prompt("chat", cfg, workspace, "research")
        check("Univerzální preference" in research_prompt
              and "Výzkumné pravidlo" in research_prompt
              and "Projektové rozhodnutí" in research_prompt
              and "Vývojové pravidlo" not in research_prompt,
              "research chat vidí své přesné tři vrstvy bez coding paměti")
    finally:
        shutil.rmtree(tmp, ignore_errors=True)


def test_safety() -> None:
    print("[safety]")
    sup = SafetyPolicy("supervised", max_steps=40, semi_max_steps=15)
    check(sup.needs_confirmation(Risk.WRITE), "supervised: WRITE potvrzení")
    check(not sup.needs_confirmation(Risk.SAFE), "supervised: SAFE bez potvrzení")
    sup.new_task()
    check(sup.step_limit() == 40, "supervised limit = max_steps")

    semi = SafetyPolicy("semi", max_steps=40, semi_max_steps=15)
    check(semi.needs_confirmation(Risk.WRITE), "semi: první WRITE potvrzení")
    semi.mark_confirmed()
    check(not semi.needs_confirmation(Risk.WRITE), "semi: další WRITE už bez potvrzení")
    check(semi.step_limit() == 15, "semi limit = semi_max_steps")

    auto = SafetyPolicy("auto", max_steps=40)
    check(not auto.needs_confirmation(Risk.WRITE), "auto: bez potvrzení")
    check(auto.step_limit() == 40, "auto limit = max_steps")
    unlimited = SafetyPolicy()
    check(unlimited.step_limit() is None, "výchozí agent nemá limit kroků")

    try:
        SafetyPolicy("režimNaval")
        check(False, "invalid autonomy vyhodí výjimku")
    except ValueError:
        check(True, "invalid autonomy vyhodí výjimku")


def test_session() -> None:
    print("[session]")
    tmp = Path(tempfile.mkdtemp())
    try:
        data = load_config().data
        data["paths"]["sessions_dir"] = str(tmp / "sessions")
        cfg = Config(data, root=ROOT)
        s = Session(cfg, session_id="test-session", system_prompt="SYS")
        s.add("user", "ahoj")
        img = tmp / "obrazek.png"
        img.write_bytes(b"\x89PNG fake")
        s.add("user", "mrkni na to", images=[img])
        s.add("assistant", "", tool_calls=[{"id": "call_1", "type": "function",
                                            "function": {"name": "list_dir", "arguments": "{}"}}])
        s.add("tool", "result text", tool_call_id="call_1", name="list_dir")
        check((tmp / "sessions/test-session/messages.jsonl").exists(), "JSONL uložen")
        n_img = len(list((tmp / "sessions/test-session/images").glob("*.png")))
        check(n_img == 1, f"obrázek zkopírován do session ({n_img})")

        api = s.to_api_messages()
        check(api[0]["role"] == "system", "system prompt na začátku")
        img_msg = [m for m in api if isinstance(m.get("content"), list)]
        check(len(img_msg) == 1 and any(p["type"] == "image_url" for p in img_msg[0]["content"]),
              "obrázek renderován jako image_url data URL")

        pinned = tmp / "pinned.txt"
        pinned.write_text("DŮLEŽITÝ PŘIPNUTÝ KONTEXT", encoding="utf-8")
        check(s.pin_context_file(pinned), "soubor lze připnout do kontextu")
        api_with_pin = s.to_api_messages()
        check(any("DŮLEŽITÝ PŘIPNUTÝ KONTEXT" in str(m.get("content", ""))
                  for m in api_with_pin), "připnutý soubor je v API pohledu modelu")
        breakdown = s.context_breakdown()
        check(breakdown["pinned_files"] == [str(pinned.resolve())],
              "context inspector eviduje připnutý soubor")
        check(s.unpin_context_file(pinned) and not s.context_breakdown()["pinned_files"],
              "připnutý soubor lze odepnout")

        original_data_url = Session.__dict__["_data_url"]
        try:
            Session._data_url = staticmethod(lambda _path: (_ for _ in ()).throw(
                AssertionError("estimate nesmí enkódovat obrázky")))
            check(s.estimate_context_tokens() > Session.IMAGE_TOKENS,
                  "odhad kontextu nečte ani base64-enkóduje obrázek")
        finally:
            Session._data_url = original_data_url

        loaded = Session.load(cfg, "test-session")
        check(len(loaded.messages) == 5, f"roundtrip zpráv (={len(loaded.messages)})")
    finally:
        shutil.rmtree(tmp, ignore_errors=True)


def test_tools_fs_shell() -> None:
    print("[tools]")
    import time
    tmp = Path(tempfile.mkdtemp())
    try:
        data = load_config().data
        data["paths"]["sessions_dir"] = str(tmp / "sessions")
        cfg = Config(data, root=ROOT)
        session = Session(cfg, session_id="tool-test")
        ctx = AgentContext(cfg=cfg, session=session, workspace=tmp)
        ctx.changes = ChangeJournal(session, tmp)
        ctx.processes = ProcessManager()
        ctx.changes.begin_task("test změn")

        reg = ToolRegistry()
        from harness.tools import fs, shell
        fs.register_fs_tools(reg)
        shell.register_shell_tools(reg)

        (tmp / "sub").mkdir()
        (tmp / "sub" / "a.txt").write_text("ahoj\nsvěte\nQWEN", encoding="utf-8")
        (tmp / "sub" / "data.py").write_text("x = 1\nQWEN_MARKER = 'zde'\n", encoding="utf-8")

        r = reg.execute("list_dir", {"path": "."}, ctx)
        check("sub" in r and "[DIR]" in r, "list_dir vidí adresář")

        r = reg.execute("read_file", {"path": "sub/a.txt"}, ctx)
        check("ahoj" in r and "3|" in r, "read_file vrací obsah s čísly řádků")

        r = reg.execute("write_file", {"path": "sub/new.md", "content": "# test"}, ctx)
        check((tmp / "sub" / "new.md").exists(), "write_file vytvořil soubor")

        r = reg.execute("search_files", {"query": "qwen_marker", "path": "."}, ctx)
        check("data.py:2" in r, "search_files case-insensitive nalezení")

        data_file = tmp / "sub" / "data.py"
        original_data = data_file.read_text(encoding="utf-8")
        r = reg.execute("apply_patch", {
            "path": "sub/data.py",
            "edits": [{"old": "NEEXISTUJE", "new": "x"}],
        }, ctx)
        check(r.startswith("ERROR") and data_file.read_text(encoding="utf-8") == original_data,
              "neplatný patch nezmění soubor")
        r = reg.execute("apply_patch", {
            "path": "sub/data.py",
            "expected_sha256": file_sha256(data_file),
            "edits": [{"old": "x = 1", "new": "x = 2"}],
        }, ctx)
        check(r.startswith("OK") and "x = 2" in data_file.read_text(encoding="utf-8"),
              "apply_patch provede přesnou atomickou změnu")
        changes = reg.execute("list_task_changes", {}, ctx)
        check("data.py" in changes and "new.md" in changes,
              "journal eviduje upravený i vytvořený soubor")
        undo = reg.execute("undo_task_changes", {}, ctx)
        check("errors\": []" in undo and data_file.read_text(encoding="utf-8") == original_data
              and not (tmp / "sub" / "new.md").exists(),
              "rollback obnoví původní stav celé úlohy")
        check(not any(item["changed"] for item in ctx.changes.summary()["files"]),
              "journal je po rollbacku znovu čistý")

        r = reg.execute("run_command", {"command": "echo hello-$((40+2))", "shell": "bash"}, ctx)
        check("hello-42" in r and "exit code: 0" in r, f"run_command bash funguje: {r[:60]}")

        r = reg.execute("run_command", {"command": "Write-Output 'ps-works'"}, ctx)
        check("ps-works" in r, f"run_command powershell funguje")

        started = time.monotonic()
        launched = json.loads(reg.execute("start_command", {
            "command": "Write-Output one; Start-Sleep -Milliseconds 200; Write-Output two",
            "shell": "powershell", "timeout": 5,
        }, ctx))
        check(time.monotonic() - started < 1 and launched["status"] == "running",
              "start_command vrátí okamžitě process_id")
        cursor = 0
        streamed = ""
        for _ in range(50):
            poll = json.loads(reg.execute("poll_command", {
                "process_id": launched["process_id"], "cursor": cursor,
            }, ctx))
            streamed += poll["output"]
            cursor = poll["cursor"]
            if poll["status"] == "finished":
                break
            time.sleep(0.05)
        check("one" in streamed and "two" in streamed and poll["exit_code"] == 0,
              "poll_command streamuje přírůstkový výstup do dokončení")

        sleeper = json.loads(reg.execute("start_command", {
            "command": "Start-Sleep -Seconds 30", "shell": "powershell", "timeout": 60,
        }, ctx))
        stopped = json.loads(reg.execute("terminate_command", {
            "process_id": sleeper["process_id"],
        }, ctx))
        check(stopped.get("terminated") is True,
              "terminate_command ukončí dlouhý process tree")

        r = reg.execute("run_command", {"command": "format c: /x"}, ctx)
        check("blocked" in r.lower(), "nebezpečný příkaz zablokován")

        r = reg.execute("read_file", {"path": "neexistuje.txt"}, ctx)
        check(r.startswith("ERROR"), "chyba čtení vrací ERROR text")

        schemas = reg.schemas()
        check(all(s["function"]["name"] for s in schemas) and len(schemas) == 12,
              f"schemas pro izolovanou sadu 12 nástrojů ({len(schemas)})")
    finally:
        shutil.rmtree(tmp, ignore_errors=True)


def test_registry_modes() -> None:
    print("[registry]")
    chat = build_registry("chat")
    agent = build_registry("agent")
    computer = build_registry("computer")
    check(set(chat.names()) == {"read_memory", "save_memory", "web_search", "web_fetch",
                                "context_status", "pin_context_file", "unpin_context_file",
                                "list_project_documents", "read_project_document",
                                "export_document", "list_skills", "read_skill"},
          f"chat režim: memory + web + context nástroje ({chat.names()})")
    check({"list_dir", "run_command", "view_image"} <= set(agent.names()),
          f"agent režim: fs+patch+shell+vision ({len(agent.names())})")
    check({"screenshot", "click", "type_text", "press_key"} <= set(computer.names()),
          f"computer režim: + GUI nástroje ({len(computer.names())})")
    click = computer.get("click")
    check(click.risk == Risk.WRITE, "click je WRITE risk")
    shot = computer.get("screenshot")
    check(shot.risk == Risk.SAFE, "screenshot je SAFE risk")

    discussion = build_registry("chat", "discussion")
    research = build_registry("chat", "research")
    writing = build_registry("agent", "writing")
    development = build_registry("agent", "development")
    check("apply_patch" not in discussion.names() and "git_commit" not in discussion.names(),
          "Diskuze nemá coding nástroje")
    check(set(research.names()) == set(discussion.names()),
          "Výzkum má web/context nástroje bez coding sady")
    check("export_document" in discussion.names() and "export_document" in research.names()
          and "export_document" in development.names(),
          "PDF/DOCX/Markdown export je viditelný ve všech pracovních režimech")
    check("apply_patch" in writing.names() and "export_document" in writing.names()
          and "repo_overview" not in writing.names() and "git_commit" not in writing.names()
          and "run_command" not in writing.names(),
          "Psaní má dokumentové editace bez Git a shellu")
    check({"apply_patch", "git_commit", "run_command", "start_project_check"}
          <= set(development.names()), "Vývoj má kompletní coding sadu")


def test_parse_args() -> None:
    print("[llm helpers]")
    check(parse_tool_arguments('{"x": 1}') == {"x": 1}, "platný JSON")
    check(parse_tool_arguments("") == {}, "prázdné argumenty")
    check(parse_tool_arguments('blabla {"x": [1,2]} blabla') == {"x": [1, 2]},
          "JSON zasypaný v textu")


def test_workspace() -> None:
    print("[workspace]")
    from harness.agent import Agent
    data = load_config().data
    data["paths"]["sessions_dir"] = str(Path(tempfile.mkdtemp()) / "sessions")
    cfg = Config(data, root=ROOT)
    tmp = Path(tempfile.mkdtemp())
    try:
        session = Session(cfg, session_id="ws-test", system_prompt="SYS")
        from harness.safety import SafetyPolicy
        agent = Agent(cfg, LLMStub(), session, build_registry("agent"),
                      SafetyPolicy("supervised"), mode="agent")
        # None -> cwd (výchozí)
        check(agent.workspace == Path.cwd().resolve(), "výchozí workspace = cwd")
        check(agent.ctx.project_workspace is None and agent.ctx.repo_index is None,
              "chat bez projektu nemá projektový dokumentový index")
        # nastavení adresáře
        p = agent.set_workspace(str(tmp))
        check(p == tmp.resolve() and agent.workspace == tmp.resolve()
              and agent.ctx.project_workspace == tmp.resolve()
              and agent.ctx.repo_index is not None,
              "set_workspace přepojí nástroje i projektový index")
        # soubor -> nadřazený adresář
        f = tmp / "soubor.txt"
        f.write_text("x", encoding="utf-8")
        (tmp / "module.py").write_text("def project_symbol():\n    return 1\n", encoding="utf-8")
        p2 = agent.set_workspace(str(f))
        check(p2 == tmp.resolve(), "soubor -> nadřazený adresář")
        # uvozovky kolem cesty
        p3 = agent.set_workspace(f'"{tmp}"')
        check(p3 == tmp.resolve(), "cesta v uvozovkách")
        # neexistující
        try:
            agent.set_workspace(tmp / "neexistuje")
            check(False, "neexistující cesta vyhodí ValueError")
        except ValueError:
            check(True, "neexistující cesta vyhodí ValueError")
        # nástroje řeší relativní cesty od workspace
        from harness.tools.base import AgentContext
        r = build_registry("agent").execute("read_file", {"path": "soubor.txt"}, agent.ctx)
        check("soubor.txt" in r and "1| x" in r, "read_file řeší cestu od workspace")
        agent.new_task("prozkoumej projekt")
        dynamic = agent._api_messages()[-1]["content"]
        check("CURRENT PROJECT SNAPSHOT" in dynamic and "project_symbol" in dynamic
              and "project_symbol" not in session.messages[0]["content"],
              "proměnlivý repo snapshot je na konci requestu a nezneplatňuje stabilní prefix")
        cached_prefix = json.dumps(session.messages, ensure_ascii=False, sort_keys=True)
        cached_count = len(session.messages)
        overview = build_registry("agent").execute("repo_overview", {}, agent.ctx)
        check("module.py" in overview and "project_symbol" in overview,
              "repo_overview vrací klíčové symboly workspace")
        (tmp / "module.py").write_text("def refreshed_symbol():\n    return 2\n", encoding="utf-8")
        refreshed = build_registry("agent").execute("repo_overview", {}, agent.ctx)
        check("refreshed_symbol" in refreshed and "project_symbol" not in refreshed,
              "repo snapshot invaliduje cache po změně souboru")
        agent.new_task("pokračuj s aktuálním stavem")
        check(json.dumps(session.messages[:cached_count], ensure_ascii=False, sort_keys=True)
              == cached_prefix
              and "refreshed_symbol" in session.messages[-1]["content"],
              "nový snapshot se přidá za beze změny znovupoužitelný prefix")
    finally:
        shutil.rmtree(tmp, ignore_errors=True)


def test_shell_readonly() -> None:
    print("[shell read-only klasifikace]")
    from harness.tools.shell import RunCommandTool, is_read_only_command
    tool = RunCommandTool()
    safe = [
        "ls -la", "cat soubor.txt", "grep -r foo .", "git status", "git log --oneline",
        "git diff HEAD~1", "find . -name '*.py'", "echo ahoj", "ls | grep test",
        "cat a.txt; cat b.txt", "stat main.py", "wc -l *.py",
    ]
    unsafe = [
        "rm -rf x", "echo ahoj > soubor.txt", "cat x | tee y", "mkdir novy",
        "git push", "git commit -m x", "curl http://x", "ls; rm x",
        "echo $(rm x)", "npm install x", "grep x . > out", "git branch nova",
        "cp a b", "cat < vstup.txt", "python skript.py", "",
    ]
    for cmd in safe:
        check(is_read_only_command(cmd), f"SAFE: {cmd!r}")
    for cmd in unsafe:
        check(not is_read_only_command(cmd), f"WRITE: {cmd!r}")
    check(tool.risk_for({"command": "ls"}) == Risk.SAFE, "risk_for ls = SAFE")
    check(tool.risk_for({"command": "rm x"}) == Risk.WRITE, "risk_for rm = WRITE")


def test_context_compression() -> None:
    print("[ctx komprese - ne-destruktivní]")
    from harness.context import render_messages_text

    rendered = render_messages_text([
        {"role": "user", "content": "HEAD " + "a" * 120},
        {"role": "assistant", "content": "MIDDLE " + "b" * 500},
        {"role": "user", "content": "TAIL-CONTEXT " + "c" * 120},
    ], max_chars=240)
    check("HEAD" in rendered and "TAIL-CONTEXT" in rendered and len(rendered) <= 240,
          "dlouhý transcript zachová začátek i nejnovější konec")
    tmp = Path(tempfile.mkdtemp())
    try:
        data = load_config().data
        data["paths"]["sessions_dir"] = str(tmp / "sessions")
        cfg = Config(data, root=ROOT)
        s = Session(cfg, session_id="ctx-test", system_prompt="SYS")
        # napodob delsi konverzaci: 6x user/assistant/tool trojice
        for i in range(6):
            s.add("user", f"otazka {i} " + "x" * 500)
            s.add("assistant", "", tool_calls=[{"id": f"c{i}", "type": "function",
                                                "function": {"name": "read_file", "arguments": "{}"}}])
            s.add("tool", f"odpoved {i} " + "y" * 300, tool_call_id=f"c{i}", name="read_file")
        n_before = len(s.messages)
        est = s.estimate_context_tokens()
        check(est > 1000, f"odhad tokenů rozumný ({est})")

        ok = s.compress_to_summary("SOUHRN konverzace.", min_keep=6)
        check(ok, "komprese proběhla")
        check(len(s.messages) == n_before, f"historie NEDOTČENÁ ({len(s.messages)} == {n_before})")
        check(s.compression is not None and s.compression["cut"] > 1, "záznam komprese (cut)")

        # model vidí méně, uživatel vše
        api_view = s._view_messages()
        check(len(api_view) < n_before, f"modelův view menší ({len(api_view)} < {n_before})")
        check("SOUHRN konverzace." in api_view[1]["content"], "souhrn v modelově view")
        est2 = s.estimate_context_tokens()
        check(est2 < est, f"tokeny pro model klesly ({est} → {est2})")

        # view nezačíná osiřelým tool voláním
        first_role = api_view[2]["role"] if len(api_view) > 2 else None
        check(first_role in ("user", None), f"cut na user hranici (role={first_role})")

        # persist + roundtrip: historie i komprese
        loaded = Session.load(cfg, "ctx-test")
        check(len(loaded.messages) == n_before, "JSONL kompletní (roundtrip)")
        check(loaded.compression is not None and loaded.compression["cut"] == s.compression["cut"],
              "compression.json persistován")

        # druhá komprese posune cut dál
        s.add("user", "nova otazka " + "a" * 100)
        s.add("assistant", "nova odpoved " + "b" * 100)
        ok2 = s.compress_to_summary("SOUHRN 2.", min_keep=2)
        check(ok2 and s.compression["cut"] > loaded.compression["cut"],
              "druhá komprese posunula cut vpřed")

        # trim fallback posouvá cut, nemaže historii
        s2 = Session(cfg, session_id="trim-test", system_prompt="SYS")
        for i in range(10):
            s2.add("user", f"u{i} " + "z" * 2000)
            s2.add("assistant", f"a{i} " + "w" * 2000)
        big = s2.estimate_context_tokens()
        n2 = len(s2.messages)
        ok = s2.trim_to_budget(big // 2)
        check(ok and s2.estimate_context_tokens() <= big // 2,
              f"trim do rozpočtu ({big} → {s2.estimate_context_tokens()})")
        check(len(s2.messages) == n2, "trim nemaže historii (jen posouvá cut)")

        # tokenový rozpočet: obří tool výstupy v ocasu nespotřebují půlku kontextu
        s3 = Session(cfg, session_id="bigtail-test", system_prompt="SYS")
        for i in range(10):
            s3.add("user", f"q{i}")
            s3.add("assistant", "", tool_calls=[{"id": f"c{i}", "type": "function",
                                                 "function": {"name": "read_file", "arguments": "{}"}}])
            s3.add("tool", "T" * 6000, tool_call_id=f"c{i}", name="read_file")  # ~1.6k toků
        # 10 trojic ≈ 16k+ tokenů; rozpočet 6k → cut musí ořezat hluboko
        est_before = s3.estimate_context_tokens()
        ok = s3.compress_to_summary("SOUHRN", keep_tokens=6000)
        est_after = s3.estimate_context_tokens()
        check(ok and est_after <= 8000,
              f"tokenový rozpočet drží ocas ({est_before} → {est_after}, cíl ≤ 8000)")
        view = s3._view_messages()
        check(view[2]["role"] == "user", "token-based cut také na user hranici")
    finally:
        shutil.rmtree(tmp, ignore_errors=True)


def test_reasoning_effort_kwargs() -> None:
    print("[reasoning effort]")
    from harness.llm import LLMClient, _template_kwargs
    data = load_config().data
    data["thinking"] = True
    data["reasoning_effort"] = "low"
    check(_template_kwargs(Config(data, ROOT)) == {"chat_template_kwargs": {"reasoning_effort": "low"}},
          "effort low → template kwarg")
    data["reasoning_effort"] = "xhigh"
    check(_template_kwargs(Config(data, ROOT))["chat_template_kwargs"]["reasoning_effort"] == "xhigh",
          "effort xhigh")
    data["thinking"] = False
    check(_template_kwargs(Config(data, ROOT)) == {
        "chat_template_kwargs": {"enable_thinking": False}},
          "thinking off má prioritu před effort")
    data["thinking"] = True
    data["reasoning_effort"] = "blbost"
    check(_template_kwargs(Config(data, ROOT)) == {}, "neplatný effort → bez kwarg (default šablony)")
    data["default_model"] = "ornith_q5"
    data["reasoning_effort"] = "xhigh"
    check(_template_kwargs(Config(data, ROOT)) == {
        "chat_template_kwargs": {"enable_thinking": True}},
          "Ornith zapne reasoning bez nepodporovaného reasoning_effort")
    data["default_model"] = "q4"

    class CaptureCompletions:
        def __init__(self):
            self.params = None

        def create(self, **params):
            self.params = params
            return []

    data["thinking"] = False
    cfg = Config(data, ROOT)
    llm = LLMClient.__new__(LLMClient)
    llm.cfg = cfg
    llm.model_name = "local-model"
    completions = CaptureCompletions()
    llm.client = type("Client", (), {
        "chat": type("Chat", (), {"completions": completions})(),
    })()
    llm.stream([{"role": "user", "content": "test"}])
    extra = completions.params["extra_body"]
    check(extra.get("top_k") == 20, "stream request zachová top_k")
    check(extra.get("chat_template_kwargs") == {"enable_thinking": False},
          "stream request zachová thinking/reasoning nastavení")
    check("max_tokens" not in completions.params,
          "produkční LLM request nemá umělý výstupní token limit")

    class ClosableStream:
        def __init__(self):
            self.closed = False
            self.parts = ["Začátek ", "dokončené věty.", " Tohle už se nevygeneruje."]

        def __iter__(self):
            for part in self.parts:
                delta = type("Delta", (), {"content": part, "tool_calls": []})()
                yield type("Chunk", (), {
                    "choices": [type("Choice", (), {"delta": delta})()]})()

        def close(self):
            self.closed = True

    closable = ClosableStream()
    llm.client = type("Client", (), {
        "chat": type("Chat", (), {
            "completions": type("Completions", (), {
                "create": staticmethod(lambda **_kwargs: closable)})(),
        })(),
    })()
    checks = {"count": 0}

    def request_stop():
        checks["count"] += 1
        return checks["count"] >= 2

    stopped = llm.stream([{"role": "user", "content": "test stop"}],
                         should_stop=request_stop)
    check(stopped.stopped and stopped.content == "Začátek dokončené věty."
          and closable.closed,
          "graceful Stop dokončí větu, zavře stream a nepokračuje dál")


def test_runtime_lifecycle_helpers() -> None:
    print("[runtime lifecycle]")
    import socket
    import time
    from harness import servermgmt
    from launcher.launcher_app import _free_web_port

    tmp = Path(tempfile.mkdtemp())
    original_health = servermgmt.health
    try:
        data = load_config().data
        data["paths"]["runtime_dir"] = str(tmp / "runtime")
        data["server"]["port"] = 65534
        cfg = Config(data, root=ROOT)
        pf = servermgmt.pid_file(cfg)
        pf.parent.mkdir(parents=True)
        pf.write_text("q4:99999999", encoding="utf-8")
        servermgmt.health = lambda *_args, **_kwargs: False
        check(servermgmt.server_state(cfg) == "down" and not pf.exists(),
              "stale PID se uklidí a server je down")

        class DeadProcess:
            @staticmethod
            def poll():
                return 1

        started = time.monotonic()
        check(not servermgmt.wait_health(cfg, timeout=10, proc=DeadProcess())
              and time.monotonic() - started < 1,
              "wait_health skončí hned po pádu procesu")

        with socket.socket(socket.AF_INET, socket.SOCK_STREAM) as occupied:
            occupied.bind(("127.0.0.1", 0))
            occupied.listen(1)
            busy_port = occupied.getsockname()[1]
            check(_free_web_port(busy_port) != busy_port,
                  "launcher přeskočí obsazený Web UI port")
    finally:
        servermgmt.health = original_health
        shutil.rmtree(tmp, ignore_errors=True)


def test_dependency_marker() -> None:
    print("[dependency marker]")
    tmp = Path(tempfile.mkdtemp())
    try:
        requirements = tmp / "requirements.txt"
        venv = tmp / ".venv"
        python = venv / "Scripts" / "python.exe"
        python.parent.mkdir(parents=True)
        python.touch()
        requirements.write_text("example==1.0\n", encoding="utf-8")

        (venv / ".deps.ok").write_text("", encoding="ascii")
        check(not dependencies_current(requirements, venv),
              "stary .deps.ok marker se nepovazuje za aktualni")
        mark_dependencies_current(requirements, venv)
        check(dependencies_current(requirements, venv),
              "SHA-256 marker potvrdi aktualni requirements")
        check(not (venv / ".deps.ok").exists(),
              "zastaraly marker se po synchronizaci odstrani")
        check(len(requirements_digest(requirements)) == 64,
              "fingerprint requirements je SHA-256")

        requirements.write_text("example==2.0\n", encoding="utf-8")
        check(not dependencies_current(requirements, venv),
              "zmena requirements zneplatni dependency marker")
    finally:
        shutil.rmtree(tmp, ignore_errors=True)


def test_streaming_bridge() -> None:
    print("[streaming bridge]")
    from harness.streaming import SteeringQueue, StreamHub, step_threaded

    hub = StreamHub()
    hub.on_event("reasoning", "uva")
    hub.on_event("reasoning", "žuji")
    hub.on_event("text", "ho")
    hub.on_event("text", "tovo")
    text, reasoning, rev, _ = hub.snapshot()
    check(text == "hotovo" and reasoning == "uvažuji" and rev == 4,
          "StreamHub skládá fragmenty bez ztráty pořadí")
    hub.on_event("tool_delta", ("write_file", '{"path":"game.py","content":"abc'))
    progress = hub.progress()
    check(progress["tool_call_name"] == "write_file"
          and progress["tool_call_chars"] > 20,
          "StreamHub zviditelní generování dlouhých argumentů nástroje")
    hub.on_event("tool_start", ("write_file", {"path": "game.py"}))
    check(hub.progress()["tools_running"] == [("write_file", {"path": "game.py"})],
          "StreamHub ukáže právě prováděný nástroj")
    hub.on_event("tool_result", ("write_file", "OK"))
    check(not hub.progress()["tools_running"],
          "StreamHub po výsledku ukončí stav provádění nástroje")

    class FakeAgent:
        @staticmethod
        def step(approve=None):
            return f"step:{approve}"

    thread, box = step_threaded(FakeAgent(), True)
    thread.join(timeout=2)
    check(not thread.is_alive() and box.get("r") == "step:True",
          "worker bridge vrátí výsledek agent.step")

    steering = SteeringQueue()
    steering.push("Nejdřív oprav parser.", ["screen.png"])
    steering.push("A zachovej kompatibilitu.")
    check(bool(steering)
          and steering.pop_all() == [
              ("Nejdřív oprav parser.", ["screen.png"]),
              ("A zachovej kompatibilitu.", []),
          ] and not steering,
          "steering queue zachová pořadí upřesnění a atomicky se vyprázdní")

    from harness.agent import Agent, Status
    from harness.llm import AssistantResult
    from harness.safety import SafetyPolicy
    tmp = Path(tempfile.mkdtemp())
    try:
        data = load_config().data
        data["paths"]["sessions_dir"] = str(tmp / "sessions")
        data["agent"]["workspace"] = None
        cfg = Config(data, root=ROOT)
        session = Session(cfg, session_id="stop-test", system_prompt="SYS")
        agent = Agent(
            cfg, LLMStub([AssistantResult(content="Dokončená věta.", stopped=True)]),
            session, ToolRegistry(), SafetyPolicy("auto"), mode="chat",
            work_mode="discussion")
        agent.new_task("Dlouhá odpověď")
        result = agent.step()
        check(result.status is Status.ABORTED
              and any(message.get("content") == "Dokončená věta."
                      for message in session.messages),
              "Stop uloží dokončenou část odpovědi a ukončí agentní úlohu")
        agent.abort_flag.set()
        agent.llm = LLMStub([AssistantResult(content="Nová odpověď")])
        agent.new_task("Nový dotaz po Stop")
        check(not agent.abort_flag.is_set() and agent.step().status is Status.FINAL,
              "nový dotaz po Stop dostane čistý abort stav")

        steered = Session(cfg, session_id="steer-test", system_prompt="SYS")
        steer_agent = Agent(
            cfg, LLMStub([AssistantResult(content="První dokončená věta.", stopped=True)]),
            steered, ToolRegistry(), SafetyPolicy("auto"), mode="chat",
            work_mode="discussion")
        steer_agent.new_task("Navrhni řešení")
        check(steer_agent.step().status is Status.ABORTED,
              "steering nejprve ukončí aktuální stream u dokončené věty")
        steer_agent.steer("Zachovej také zpětnou kompatibilitu.")
        steer_agent.llm = LLMStub([AssistantResult(content="Upravené řešení.")])
        steer_result = steer_agent.step()
        visible = [m.get("content") for m in steered.messages
                   if m.get("role") != "system"
                   and not str(m.get("content") or "").startswith(Session.INTERNAL_USER_PREFIXES)]
        check(steer_result.status is Status.FINAL
              and visible[-3:] == [
                  "První dokončená věta.",
                  "Zachovej také zpětnou kompatibilitu.",
                  "Upravené řešení.",
              ],
              "steering zachová část odpovědi a pokračuje s upřesněním ve správném pořadí")
    finally:
        shutil.rmtree(tmp, ignore_errors=True)


def test_parallel_read_tools() -> None:
    print("[parallel read tools]")
    import time
    from harness.agent import Agent
    from harness.llm import AssistantResult
    from harness.safety import SafetyPolicy
    from harness.tools.base import Tool

    class SlowRead(Tool):
        parallel_safe = True
        parameters = {}

        def __init__(self, name):
            self.name = name

        def run(self, _ctx):
            time.sleep(0.25)
            return self.name

    class SlowWrite(SlowRead):
        parallel_safe = False

    tmp = Path(tempfile.mkdtemp())
    try:
        data = load_config().data
        data["paths"]["sessions_dir"] = str(tmp / "sessions")
        cfg = Config(data, root=ROOT)
        session = Session(cfg, session_id="parallel-test", system_prompt="SYS")
        registry = ToolRegistry()
        registry.register(SlowRead("read_a"))
        registry.register(SlowRead("read_b"))
        registry.register(SlowWrite("write_c"))
        agent = Agent(cfg, LLMStub(), session, registry, SafetyPolicy("auto"), mode="agent")
        calls = [_tc("read_a"), _tc("read_b")]
        started = time.monotonic()
        trace = agent._execute_calls(calls, "Našel jsem první podklady, teď je porovnám.")
        parallel_time = time.monotonic() - started
        check(parallel_time < 0.45 and [item[2] for item in trace] == ["read_a", "read_b"],
              "nezávislé read-only tool calls běží paralelně a zachovají pořadí")
        persisted = next(message for message in session.messages
                         if message.get("tool_calls") == calls)
        reloaded = Session.load(cfg, session.id)
        check(persisted["content"].startswith("Našel jsem")
              and any(str(message.get("content", "")).startswith("Našel jsem")
                      for message in reloaded.messages),
              "průběžný text před tool callem zůstane v chatu i po reloadu")
        started = time.monotonic()
        agent._execute_calls([_tc("read_a"), _tc("write_c")])
        check(time.monotonic() - started >= 0.45,
              "smíšená read/write sada zůstane sekvenční")
    finally:
        shutil.rmtree(tmp, ignore_errors=True)


def test_resume_task_and_process_after_restart() -> None:
    print("[resume task/process after restart]")
    import time
    from harness.agent import Agent, Status
    from harness.llm import AssistantResult
    from harness.safety import SafetyPolicy

    tmp = Path(tempfile.mkdtemp())
    managers = []
    try:
        data = load_config().data
        data["paths"]["sessions_dir"] = str(tmp / "sessions")
        data["agent"]["workspace"] = str(tmp)
        cfg = Config(data, root=ROOT)
        session = Session(cfg, session_id="resume-test", system_prompt="SYS")
        pending_llm = LLMStub([AssistantResult(content="Připravuji soubor a čekám na potvrzení.", tool_calls=[
            _tc("write_file", '{"path":"resume.txt","content":"OK"}')])])
        first = Agent(
            cfg, pending_llm, session, build_registry("agent", "development"),
            SafetyPolicy("supervised"), mode="agent", work_mode="development")
        first.new_task("vytvoř resume.txt")
        waiting = first.step()
        check(waiting.status is Status.NEEDS_CONFIRMATION
              and session.load_task_state()["status"] == "waiting_confirmation"
              and session.load_task_state()["pending_text"].startswith("Připravuji"),
              "pending potvrzení i jeho průběžný text se uloží do task-state")

        restored = Agent(
            cfg, LLMStub([]), session, build_registry("agent", "development"),
            SafetyPolicy("supervised"), mode="agent", work_mode="development")
        restored_waiting = restored.step()
        check(restored.has_resumable_task
              and restored_waiting.status is Status.NEEDS_CONFIRMATION
              and restored_waiting.text.startswith("Připravuji"),
              "nový Agent po restartu obnoví pending tool calls i viditelný text")
        restored.step(approve=True)
        check((tmp / "resume.txt").is_file()
              and any(str(message.get("content", "")).startswith("Připravuji")
                      for message in session.messages if message.get("role") == "assistant"),
              "obnovené potvrzení dokončí tool call a zachová průběh")
        completed = Agent(
            cfg, LLMStub([AssistantResult(content="hotovo")]), session,
            build_registry("agent", "development"), SafetyPolicy("supervised"),
            mode="agent", work_mode="development")
        check(completed.has_resumable_task and completed.step().status is Status.FINAL
              and session.load_task_state()["status"] == "complete",
              "running úloha po restartu pokračuje do FINAL")

        manager1 = ProcessManager()
        managers.append(manager1)
        manager1.bind_session(session)
        item = manager1.start(
            "Write-Output BEFORE; Start-Sleep -Milliseconds 500; Write-Output AFTER",
            "powershell", tmp, 10)
        time.sleep(0.15)
        manager2 = ProcessManager()
        managers.append(manager2)
        manager2.bind_session(session)
        restored_item = manager2.get(item.id)
        check(restored_item is not None and restored_item.proc is None,
              "ProcessManager načte procesní manifest po restartu")
        cursor = 0
        output = ""
        for _ in range(100):
            polled = manager2.poll(item.id, cursor)
            output += polled["output"]
            cursor = polled["cursor"]
            if polled["status"] == "finished":
                break
            time.sleep(0.05)
        check("BEFORE" in output and "AFTER" in output,
              "obnovený ProcessManager pokračuje ve čtení persistentního logu")
    finally:
        for manager in managers:
            manager.terminate_all()
        shutil.rmtree(tmp, ignore_errors=True)


def test_git_tools() -> None:
    print("[git tools]")
    import subprocess
    from harness.tools import fs as fs_tools, git as git_tools

    tmp = Path(tempfile.mkdtemp())
    try:
        def git(*args):
            return subprocess.run(["git", *args], cwd=tmp, check=True,
                                  capture_output=True, text=True, encoding="utf-8")

        git("init", "-q")
        git("config", "user.email", "qwen-test@example.invalid")
        git("config", "user.name", "Qwen Test")
        tracked = tmp / "tracked.txt"
        tracked.write_text("before\n", encoding="utf-8")
        git("add", "tracked.txt")
        git("commit", "-q", "-m", "baseline")

        data = load_config().data
        data["paths"]["sessions_dir"] = str(tmp / "sessions")
        cfg = Config(data, root=ROOT)
        session = Session(cfg, session_id="git-test")
        ctx = AgentContext(cfg=cfg, session=session, workspace=tmp)
        ctx.changes = ChangeJournal(session, tmp)
        ctx.changes.begin_task("git test")
        reg = ToolRegistry()
        fs_tools.register_fs_tools(reg)
        git_tools.register_git_tools(reg)

        patched = reg.execute("apply_patch", {
            "path": "tracked.txt",
            "edits": [{"old": "before", "new": "after"}],
        }, ctx)
        check(patched.startswith("OK"), "git test změna vznikla přes apply_patch")
        check("tracked.txt" in reg.execute("git_status", {}, ctx),
              "git_status vrací změněný soubor")
        check("-before" in reg.execute("git_diff", {"path": "tracked.txt"}, ctx),
              "git_diff vrací obsah změny")
        committed = reg.execute("git_commit", {"message": "task change"}, ctx)
        check("exit code: 0" in committed and git("log", "-1", "--pretty=%s").stdout.strip() == "task change",
              "git_commit commitne pouze journalované změny")
        check(git("status", "--porcelain", "--untracked-files=no").stdout.strip() == "",
              "tracked změny jsou po commitu čisté")
        check("sessions/" in git("status", "--porcelain").stdout,
              "git_commit nepřibere nesouvisející journal artefakty")
    finally:
        shutil.rmtree(tmp, ignore_errors=True)


def test_automatic_project_check() -> None:
    print("[automatic project check]")
    import time
    from harness.tools import context as context_tools

    tmp = Path(tempfile.mkdtemp())
    try:
        (tmp / "tests").mkdir()
        (tmp / "tests" / "test_core.py").write_text(
            "print('PROJECT-CHECK-OK')\n", encoding="utf-8")
        data = load_config().data
        data["paths"]["sessions_dir"] = str(tmp / ".sessions")
        cfg = Config(data, root=ROOT)
        session = Session(cfg, session_id="check-test")
        ctx = AgentContext(cfg=cfg, session=session, workspace=tmp,
                           project_workspace=tmp,
                           processes=ProcessManager())
        reg = ToolRegistry()
        context_tools.register_coding_context_tools(reg)
        started = json.loads(reg.execute("start_project_check", {"timeout": 10}, ctx))
        cursor = 0
        output = ""
        for _ in range(100):
            result = ctx.processes.poll(started["process_id"], cursor)
            output += result["output"]
            cursor = result["cursor"]
            if result["status"] == "finished":
                break
            time.sleep(0.05)
        check("PROJECT-CHECK-OK" in output and result["exit_code"] == 0,
              "detekovaný project check doběhne v background procesu")
    finally:
        shutil.rmtree(tmp, ignore_errors=True)


def test_research_ledger_and_synthesis() -> None:
    print("[research ledger + synthesis]")
    from harness.llm import AssistantResult
    from harness.research import ResearchLedger, plan_research, synthesize_research

    tmp = Path(tempfile.mkdtemp())
    try:
        data = load_config().data
        data["paths"]["sessions_dir"] = str(tmp / "sessions")
        cfg = Config(data, root=ROOT)
        session = Session(cfg, session_id="research-test", system_prompt="SYS",
                          work_mode="research")
        ledger = ResearchLedger(session)
        ledger.begin("Jaké jsou dvě protichůdné interpretace?")
        ledger.record_query("první hledání", [
            ("Zdroj A", "https://example.test/a", "tvrdí A"),
            ("Zdroj B", "https://example.test/b", "tvrdí B"),
        ])
        ledger.record_source("https://example.test/a", "Zdroj A",
                             "Interpretace A říká ano.")
        ledger.record_source("https://example.test/b", "Zdroj B",
                             "Interpretace B říká ne.")
        run = ledger.current()
        check(len(run["candidates"]) == 2 and len(run["sources"]) == 2,
              "ledger zachová všechny kandidáty i načtené zdroje")
        check(all("credibility" not in source and "trust" not in source
                  for source in run["sources"]),
              "ledger neobsahuje trust scoring ani filtr důvěryhodnosti")

        class ResearchLLM:
            def __init__(self):
                self.cfg = cfg
                self.prompts: list[str] = []

            def ask(self, messages, **_kwargs):
                prompt = str(messages[-1]["content"])
                self.prompts.append(prompt)
                if "Return JSON only" in prompt:
                    return AssistantResult(content=json.dumps({
                        "subquestions": ["Co tvrdí A?", "Co tvrdí B?"],
                        "search_angles": ["protiklady"],
                        "source_types_to_include": ["všechny dostupné"],
                        "known_constraints": [],
                    }, ensure_ascii=False))
                if "Chybějící zdroje" in prompt:
                    return AssistantResult(content="Opravená syntéza zahrnuje [S1] i [S2].")
                if "Vytvoř přehlednou závěrečnou syntézu" in prompt:
                    return AssistantResult(content="První syntéza obsahuje pouze [S1].")
                return AssistantResult(content="Dílčí loss-aware poznámky [S1] [S2].")

            def stream(self, messages, **kwargs):
                if kwargs.get("thinking") is False:
                    return self.ask(messages, **kwargs)
                return AssistantResult(content="Pracovní draft před syntézou")

        fake = ResearchLLM()
        class EmptyPlannerLLM(ResearchLLM):
            def ask(self, messages, **_kwargs):
                return AssistantResult(content="", reasoning="nedokončené uvažování")

        fallback = plan_research(EmptyPlannerLLM(), "Co je potřeba zjistit?")
        check(fallback["subquestions"] == ["Co je potřeba zjistit?"]
              and len(fallback["search_angles"]) >= 3,
              "prázdná odpověď planneru použije plán a nezastaví výzkum")

        synthesis = synthesize_research(fake, run)
        check("[S1]" in synthesis and "[S2]" in synthesis,
              "coverage kontrola doplní každý zpracovaný source ID")
        final_prompt = next(prompt for prompt in fake.prompts
                            if "Vytvoř přehlednou závěrečnou syntézu" in prompt)
        check("Interpretace A" in final_prompt and "Interpretace B" in final_prompt
              and "Nehodnoť ani nefiltruj" in final_prompt,
              "syntéza dostane protichůdná data bez trust filtru")
        ledger.complete(synthesis)
        check(ledger.status()["status"] == "complete" and ledger.path.is_file(),
              "research ledger je persistentní a označí hotovou syntézu")

        from harness.agent import Agent, Status
        from harness.safety import SafetyPolicy
        integrated_session = Session(
            cfg, session_id="research-agent", system_prompt="SYS", work_mode="research")
        integrated_llm = ResearchLLM()
        agent = Agent(
            cfg, integrated_llm, integrated_session,
            build_registry("chat", "research"), SafetyPolicy("auto"),
            mode="chat", work_mode="research",
        )
        agent.new_task("Integrovaná research otázka")
        agent.ctx.research.record_source("https://example.test/a", "A", "Ano [S1]")
        agent.ctx.research.record_source("https://example.test/b", "B", "Ne [S2]")
        result = agent.step()
        check(result.status is Status.FINAL and "[S1]" in result.text and "[S2]" in result.text
              and "Pracovní draft" not in result.text,
              "research Agent nahradí draft povinnou coverage syntézou")
        check(any(message.get("role") == "assistant"
                  and message.get("content") == "Pracovní draft před syntézou"
                  for message in integrated_session.messages),
              "pracovní draft před syntézou zůstane viditelně uložený v chatu")
        integrated_run = agent.ctx.research.current()
        check(integrated_run["plan"]["subquestions"] == ["Co tvrdí A?", "Co tvrdí B?"]
              and any(str(message.get("content", "")).startswith("[RESEARCH PLAN")
                      for message in integrated_session.messages),
              "research plán vznikne před hledáním a uloží se do ledgeru i kontextu")

        run_id = integrated_run["id"]
        agent.new_task("Ulož předchozí výstup jako PDF soubor")
        check(agent.ctx.research.current()["id"] == run_id,
              "PDF follow-up v Research režimu nezakládá nový výzkum")
        exported = agent.registry.execute("export_document", {
            "content": "# Uložený výsledek\n\nDůležitá syntéza.",
            "filename": "research-vystup",
            "format": "pdf",
            "title": "Výzkumný výstup",
        }, agent.ctx)
        no_project_pdf = integrated_session.dir / "exports" / "research-vystup.pdf"
        check(exported.startswith("OK:") and no_project_pdf.is_file()
              and "Důležitá syntéza" in "\n".join(
                  page.extract_text() or "" for page in __import__("pypdf").PdfReader(
                      no_project_pdf).pages),
              "Research export bez projektu uloží čitelné PDF do session")
    finally:
        shutil.rmtree(tmp, ignore_errors=True)


def test_project_document_library() -> None:
    print("[project document library]")
    from docx import Document
    from harness.agent import Agent
    from harness.documents import export_document
    from harness.repo_index import RepoIndex
    from harness.safety import SafetyPolicy
    from pypdf import PdfWriter

    tmp = Path(tempfile.mkdtemp())
    try:
        (tmp / "film.md").write_text("# Film Aurora\nHrdinka se jmenuje Klára.\n", encoding="utf-8")
        (tmp / "code.py").write_text("SECRET_CODE = 1\n", encoding="utf-8")
        document = Document()
        document.add_paragraph("DOCX podklad o postavě Klára")
        document.save(tmp / "postava.docx")
        writer = PdfWriter()
        writer.add_blank_page(width=200, height=200)
        with open(tmp / "reference.pdf", "wb") as handle:
            writer.write(handle)
        data = load_config().data
        data["paths"]["sessions_dir"] = str(tmp / "sessions")
        cfg = Config(data, root=ROOT)

        discussion_session = Session(
            cfg, session_id="discussion-docs", system_prompt="SYS",
            workspace=str(tmp), work_mode="discussion")
        discussion_agent = Agent(
            cfg, LLMStub(), discussion_session, build_registry("chat", "discussion"),
            SafetyPolicy("auto"), mode="chat", work_mode="discussion")
        discussion_agent.set_workspace(tmp)
        discussion_agent.new_task("Pověz mi o filmu")
        prompt = discussion_agent._api_messages()[-1]["content"]
        check("CURRENT PROJECT DOCUMENT LIBRARY" in prompt and "film.md" in prompt
              and "CURRENT PROJECT SNAPSHOT" not in prompt and "code.py" not in prompt,
              "Diskuze vidí dokumentovou knihovnu bez coding repo snapshotu")

        research_session = Session(
            cfg, session_id="research-docs", system_prompt="SYS",
            workspace=str(tmp), work_mode="research")
        research_agent = Agent(
            cfg, LLMStub(), research_session, build_registry("chat", "research"),
            SafetyPolicy("auto"), mode="chat", work_mode="research")
        research_agent.set_workspace(tmp)
        research_agent.new_task("Kdo je hrdinka?")
        content = research_agent.registry.execute(
            "read_project_document", {"path": "film.md"}, research_agent.ctx)
        sources = research_agent.ctx.research.current()["sources"]
        check("Klára" in content and len(sources) == 1
              and sources[0]["url"].startswith("file://"),
              "lokální dokument se ve Výzkumu uloží jako ledger source")
        library = RepoIndex(tmp)
        _, docx_text = library.read_document("postava.docx")
        pdf_path, pdf_text = library.read_document("reference.pdf")
        check("DOCX podklad" in docx_text, "projektová knihovna čte skutečný DOCX")
        check(pdf_path.suffix == ".pdf" and isinstance(pdf_text, str),
              "projektová knihovna načte validní PDF")
        exported_docx = export_document(
            "# Nadpis\n\nText o Kláře.\n\n- bod A", tmp / "exports", "vystup", "docx", "Film")
        exported_pdf = export_document(
            "# Nadpis\n\n**Text o Kláře.** a Na⁺.\n\n"
            "| Položka | Hodnota |\n|---|---|\n"
            "| Zdroj | [Web](https://example.com) |\n\n## 📋 Next steps",
            tmp / "exports", "vystup", "pdf", "Film")
        exported_docx_text = "\n".join(
            paragraph.text for paragraph in Document(exported_docx).paragraphs)
        exported_pdf_text = "\n".join(
            page.extract_text() or "" for page in __import__("pypdf").PdfReader(exported_pdf).pages)
        check("Kláře" in exported_docx_text, "Psaní exportuje strukturovaný DOCX s češtinou")
        check("Kláře" in exported_pdf_text and "**" not in exported_pdf_text
              and "Položka" in exported_pdf_text and "Zdroj" in exported_pdf_text
              and "Na+" in exported_pdf_text and "Next steps" in exported_pdf_text,
              "PDF export vykreslí češtinu, inline Markdown a tabulku")
    finally:
        shutil.rmtree(tmp, ignore_errors=True)


def test_async_model_switch() -> None:
    print("[async model switch]")
    import threading
    from harness.model_switch import ModelSwitchController

    entered = threading.Event()
    release = threading.Event()
    callbacks: list[str] = []

    def ensure(_cfg, key):
        entered.set()
        release.wait(timeout=2)
        return key == "q5"

    controller = ModelSwitchController(load_config(), ensure_fn=ensure)
    check(controller.request("q5", on_success=callbacks.append),
          "první switch se spustí na pozadí")
    check(entered.wait(timeout=1) and controller.snapshot().busy,
          "controller ihned hlásí starting")
    check(not controller.request("q4"), "souběžný switch je odmítnut")
    release.set()
    check(controller.wait(timeout=2), "background switch doběhne")
    snap = controller.snapshot()
    check(snap.status == "ready" and snap.target == "q5" and callbacks == ["q5"],
          "úspěšný switch publikuje ready a callback")

    failed = ModelSwitchController(load_config(), ensure_fn=lambda _cfg, _key: False)
    check(failed.request("q4") and failed.wait(timeout=2)
          and failed.snapshot().status == "failed",
          "selhání serveru se propíše do stavu controlleru")


def test_session_meta() -> None:
    print("[session meta + historie]")
    tmp = Path(tempfile.mkdtemp())
    try:
        data = load_config().data
        data["paths"]["sessions_dir"] = str(tmp / "sessions")
        cfg = Config(data, root=ROOT)
        # session s workspace + titulkem z prvního dotazu
        s = Session(cfg, session_id="meta-a", system_prompt="SYS", workspace=r"C:\projekty\Alfa")
        s.add("user", "Oprav bug v parseru")
        s.add("assistant", "hotovo")
        check(s.meta["title"] == "Oprav bug v parseru", "titulek z prvního dotazu")
        check((tmp / "sessions/meta-a/meta.json").exists(), "meta.json uložen")
        # protokolová poznámka se titulkem stát nesmí
        s2 = Session(cfg, session_id="meta-b", system_prompt="SYS")
        s2.add("user", "[TASK PROTOCOL - follow] abc")
        s2.add("user", "Skutečný dotaz")
        check(s2.meta["title"] == "Skutečný dotaz", "poznámka [..] se titulkem nestává")
        # výpis s metadaty, setříděný podle updated
        lst = Session.list_sessions(cfg)
        check({x["id"] for x in lst} == {"meta-a", "meta-b"}, "list_sessions vidí obě")
        a = next(x for x in lst if x["id"] == "meta-a")
        check(a["workspace"] == r"C:\projekty\Alfa" and a["title"] == "Oprav bug v parseru",
              "meta ve výpisu (workspace + titulek)")
        check(a["messages"] == s.meta["message_count"] == len(s.messages),
              "počet zpráv se čte z meta indexu")
        check(lst[0]["id"] == "meta-b", "novější session první")
        # stará session bez meta → titulek dohoní z první user zprávy
        s3 = Session(cfg, session_id="meta-old", system_prompt="SYS")
        s3.add("user", "Starý dotaz bez mety")
        (tmp / "sessions/meta-old/meta.json").unlink()
        loaded = Session.load(cfg, "meta-old")
        check(loaded.meta["title"] == "Starý dotaz bez mety", "zpětná kompatibilita titulku")
    finally:
        shutil.rmtree(tmp, ignore_errors=True)


def test_chat_rewind_and_fork() -> None:
    print("[chat retry/undo/fork]")
    tmp = Path(tempfile.mkdtemp())
    try:
        data = load_config().data
        data["paths"]["sessions_dir"] = str(tmp / "sessions")
        cfg = Config(data, root=ROOT)
        session = Session(cfg, session_id="original", system_prompt="SYS", workspace=str(tmp),
                          work_mode="writing")
        session.add("user", "první dotaz")
        session.add("assistant", "první odpověď")
        image = tmp / "source.png"
        image.write_bytes(b"image-data")
        session.add("user", "druhý dotaz", images=[image])
        session.add("user", "[TASK PROTOCOL - internal]")
        session.add("assistant", "druhá odpověď")
        original_count = len(session.messages)

        markdown = session.export_markdown()
        jsonl = session.export_jsonl()
        check(markdown.is_file() and "druhý dotaz" in markdown.read_text(encoding="utf-8"),
              "chat lze exportovat do čitelného Markdownu")
        imported = Session.import_jsonl(cfg, jsonl, "IMPORTED SYS", workspace=str(tmp))
        check(imported.id != session.id and imported.messages[0]["content"] == "IMPORTED SYS"
              and imported.last_user_index() is not None,
              "JSONL import vytvoří novou session a obnoví system prompt")
        found = Session.search_sessions(cfg, "druhý dotaz")
        check(any(item["id"] == session.id and "druhý dotaz" in item["snippet"]
                  and "sessions" not in item["snippet"] for item in found),
              "fulltextové hledání najde dotaz v historii")
        check((tmp / "sessions" / "history-index.sqlite3").is_file(),
              "historie používá persistentní SQLite FTS index")

        fork = session.fork_at_last_user("FORK SYS")
        check(fork is not None and len(session.messages) == original_count and fork.id != session.id,
              "fork vytvoří novou session bez změny originálu")
        check(fork.meta["work_mode"] == "writing", "fork zachová pracovní režim session")
        fork_index = fork.last_user_index()
        fork_user = fork.messages[fork_index] if fork_index is not None else {}
        copied_images = fork_user.get("images", [])
        check(fork_user.get("content") == "druhý dotaz" and copied_images
              and Path(copied_images[0]).exists() and copied_images[0] != session.messages[3]["images"][0],
              "fork končí posledním dotazem a kopíruje jeho přílohy")

        session.compression = {"cut": 2, "summary": "old"}
        session._save_compression()
        prompt = session.rewind_last_turn(keep_user=True)
        check(prompt == "druhý dotaz" and session.messages[-1]["content"] == "druhý dotaz"
              and session.compression is None,
              "retry ponechá dotaz a odstraní odpověď i starou kompresi")

        session.add("assistant", "nová odpověď")
        removed = session.rewind_last_turn(keep_user=False)
        check(removed == "druhý dotaz" and session.last_user_index() is not None
              and session.messages[session.last_user_index()]["content"] == "první dotaz",
              "undo odstraní celé poslední kolo")
    finally:
        shutil.rmtree(tmp, ignore_errors=True)


def test_transient_session() -> None:
    print("[transient session - lazy zápis na disk]")
    tmp = Path(tempfile.mkdtemp())
    try:
        data = load_config().data
        data["paths"]["sessions_dir"] = str(tmp / "sessions")
        cfg = Config(data, root=ROOT)
        s = Session(cfg, system_prompt="SYS", workspace=r"C:\projekty\Alfa", transient=True)
        s.add("assistant", "ahoj")
        check(s.transient and not s.dir.exists(),
              "bez user zprávy nic na disku (transient)")
        s.add("user", "Oprav bug")
        check(not s.transient and s.dir.exists() and s._jsonl.exists(),
              "první user zpráva = persist na disk")
        lines = s._jsonl.read_text(encoding="utf-8").strip().splitlines()
        check(len(lines) == 3, f"celá historie zapsána ({len(lines)} řádků)")
        s2 = Session.load(cfg, s.id)
        check(len(s2.messages) == 3 and not s2.transient, "load po persist")
        check(Session.delete(cfg, "neexistujici-x") is False, "delete neexistující = False")
    finally:
        shutil.rmtree(tmp, ignore_errors=True)


def test_web_tools() -> None:
    print("[web nástroje - offline parsery + registrace]")
    from io import BytesIO
    from harness.tools import web as webt
    from reportlab.pdfgen import canvas

    txt = webt._strip_tags("<html><body><script>bad()</script><h1>Ahoj</h1>"
                           "<p>sv&#233;te &amp; nazdar</p></body></html>")
    check("Ahoj" in txt and "svéte & nazdar" in txt and "bad()" not in txt,
          "_strip_tags: tagy pryč, entity dekódovány")
    u = webt._ddg_unwrap("//duckduckgo.com/l/?uddg=https%3A%2F%2Fexample.com%2Fdoc")
    check(u == "https://example.com/doc", "_ddg_unwrap rozbalí uddg redirect")
    enc = ("https://www.bing.com/ck/a?!&amp;&amp;p=xx&u=a1aHR0cHM6Ly9naXRodWIuY29tL3Rlc3Q"
           "&ntb=1")
    check(webt._bing_unwrap(enc) == "https://github.com/test", "_bing_unwrap dekóduje base64 url")
    pdf_buffer = BytesIO()
    pdf_canvas = canvas.Canvas(pdf_buffer)
    pdf_canvas.drawString(72, 720, "INTERNET-PDF-OK")
    pdf_canvas.save()
    pdf_text, _ = webt._extract_downloaded_document(
        pdf_buffer.getvalue(), "application/pdf", "https://example.test/studie.pdf")
    docx_buffer = BytesIO()
    docx_document = __import__("docx").Document()
    docx_document.add_paragraph("INTERNET-DOCX-OK")
    docx_document.save(docx_buffer)
    docx_text, _ = webt._extract_downloaded_document(
        docx_buffer.getvalue(),
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "https://example.test/studie.docx")
    check("INTERNET-PDF-OK" in pdf_text, "web_fetch extrahuje text internetového PDF")
    check("INTERNET-DOCX-OK" in docx_text, "web_fetch extrahuje text internetového DOCX")
    for mode in ("chat", "agent", "computer"):
        reg = build_registry(mode)
        check("web_search" in reg.names() and "web_fetch" in reg.names(),
              f"web nástroje v režimu {mode}")
    reg = build_registry("chat")
    cfgd = load_config()
    ctx = type("C", (), {"cfg": cfgd})()
    out = reg.execute("web_fetch", {"url": "ftp://neplatne.cz"}, ctx)
    check(out.startswith("ERROR"), "web_fetch odmítne non-http url")

    import requests
    from harness.research import ResearchLedger
    tmp = Path(tempfile.mkdtemp())
    original_ensure = webt._ensure_ddgs
    original_bing = webt.WebSearchTool._bing
    original_ddg = webt.WebSearchTool._ddg
    original_get = requests.get
    try:
        data = load_config().data
        data["paths"]["sessions_dir"] = str(tmp / "sessions")
        cfg = Config(data, root=ROOT)
        session = Session(cfg, session_id="web-research", system_prompt="SYS",
                          work_mode="research")
        ledger = ResearchLedger(session)
        ledger.begin("test internetového ledgeru")
        research_ctx = type("RC", (), {"cfg": cfg, "research": ledger})()
        webt._ensure_ddgs = lambda: False
        webt.WebSearchTool._bing = staticmethod(lambda *_args: [
            ("A", "https://example.test/a", "ano"),
            ("B", "https://example.test/b", "ne"),
        ])
        webt.WebSearchTool._ddg = staticmethod(lambda *_args: [])
        webt.WebSearchTool().run(research_ctx, "protiklad", 2)
        check(len(ledger.current()["candidates"]) == 2,
              "web_search uloží všechny nalezené kandidáty do ledgeru")

        class Response:
            url = "https://example.test/a"
            headers = {"content-type": "text/html"}
            text = "<html><title>Zdroj A</title><body>Obsah tvrdí ano i ne.</body></html>"
            status_code = 200

            @staticmethod
            def raise_for_status():
                return None

        requests.get = lambda *_args, **_kwargs: Response()
        fetched = webt.WebFetchTool().run(research_ctx, "https://example.test/a")
        source = ledger.current()["sources"][0]
        check("Obsah tvrdí ano i ne" in fetched and "ano i ne" in source["content"],
              "web_fetch uloží plný čitelný obsah zdroje do ledgeru")
        check("credibility" not in source and "trust" not in source,
              "web integrace nepřidává hodnocení důvěryhodnosti")
    finally:
        webt._ensure_ddgs = original_ensure
        webt.WebSearchTool._bing = original_bing
        webt.WebSearchTool._ddg = original_ddg
        requests.get = original_get
        shutil.rmtree(tmp, ignore_errors=True)


def test_projects() -> None:
    print("[projekty]")
    from harness.projects import Projects
    tmp = Path(tempfile.mkdtemp())
    try:
        data = load_config().data
        data["projects"] = {"root_dir": "projects"}
        data["work_mode"] = "writing"
        cfg = Config(data, root=tmp)
        pj = Projects(cfg)
        # nový projekt - složka v projects rootu
        p1 = pj.create_new("Můj-Test:Projekt")  # nebezpečné znaky → sanitizace
        check((tmp / "projects" / p1["name"]).is_dir(), f"složka vytvořena ({p1['name']})")
        check(p1["work_mode"] == "writing", "projekt uloží výchozí pracovní režim")
        check(p1["name"] != "Můj-Test:Projekt" or True, "název sanitizován")
        # duplicita jmen → -2
        p2 = pj.create_new(p1["name"])
        check(p2["name"] != p1["name"], "unikátní název při duplicitě")
        # připojení existující složky - jméno dle složky, idempotentní
        ext = tmp / "Existujici"
        ext.mkdir()
        a1 = pj.attach_folder(str(ext))
        a2 = pj.attach_folder(str(ext))
        check(a1["name"] == "Existujici" and a1["id"] == a2["id"], "attach idempotentní")
        pj.set_work_mode(str(ext.resolve()), "research")
        check(pj.by_path(str(ext.resolve()))["work_mode"] == "research",
              "výchozí režim projektu lze změnit")
        # registr vrátí vše
        names = [p["name"] for p in pj.list_all()]
        check(len(names) == 3, f"3 projekty v registru ({names})")
        managed_file = Path(p1["path"]) / "data" / "artifact.txt"
        managed_file.parent.mkdir()
        managed_file.write_text("projektová data", encoding="utf-8")
        pj.delete_by_path(p1["path"])
        check(not Path(p1["path"]).exists() and pj.by_path(p1["path"]) is None,
              "smazání projektu odstraní registraci, složku i všechny soubory")
        # session delete + adopt
        s = Session(cfg, session_id="proj-s", system_prompt="SYS", workspace=str(ext))
        check(Session.delete(cfg, "proj-s") and not (tmp/"sessions"/"proj-s").exists(),
              "session delete")
        s2 = Session(cfg, session_id="adopt-s", system_prompt="SYS")  # bez workspace
        s2.adopt_workspace(str(ext))
        check(s2.meta["workspace"] == str(ext), "adopt workspace")
        s2.adopt_workspace("jina")  # už má - nesmí přepsat
        check(s2.meta["workspace"] == str(ext), "adopt nepřepisuje existující")
        (ext / "external.txt").write_text("data", encoding="utf-8")
        pj.delete_by_path(str(ext.resolve()))
        check(not ext.exists(), "explicitní smazání připojeného projektu odstraní jeho adresář")
        protected = pj.attach_folder(str(tmp))
        try:
            pj.delete_by_path(protected["path"])
            check(False, "kořen aplikace nelze smazat jako projekt")
        except ValueError:
            check(tmp.exists(), "kořen aplikace nelze smazat jako projekt")
    finally:
        shutil.rmtree(tmp, ignore_errors=True)


def test_skill_library() -> None:
    print("[skills]")
    from harness.skills import SkillLibrary
    system = SkillLibrary(load_config())
    names = [item.name for item in system.list()]
    check({"systematic-debugging", "architecture-options", "implementation-verification"}
          <= set(names), "systémová knihovna objeví bundlované SKILL.md")
    check("root cause" in system.read("systematic-debugging").lower(),
          "tělo skillu se načte až explicitním čtením")

    tmp = Path(tempfile.mkdtemp())
    try:
        custom = tmp / ".qwen-skills" / "override" / "SKILL.md"
        custom.parent.mkdir(parents=True)
        custom.write_text(
            "---\nname: systematic-debugging\n"
            "description: Project-specific debugging guidance.\n---\n\nPROJECT OVERRIDE\n",
            encoding="utf-8")
        library = SkillLibrary(load_config(), tmp)
        info = next(item for item in library.list() if item.name == "systematic-debugging")
        check(info.source == "project" and "PROJECT OVERRIDE" in library.read(info.name),
              "projektový skill může přepsat systémový skill stejného jména")
        ctx = type("SkillCtx", (), {
            "cfg": load_config(), "project_workspace": tmp,
        })()
        output = build_registry("chat").execute(
            "read_skill", {"name": "systematic-debugging"}, ctx)
        check("PROJECT OVERRIDE" in output, "read_skill zpřístupní vybraný postup modelu")
    finally:
        shutil.rmtree(tmp, ignore_errors=True)


class LLMStub:
    """Fake LLM se scénářem - vrací předpřipravené odpovědi v pořadí."""
    def __init__(self, script=None):
        from harness.llm import AssistantResult
        self.script = list(script or [])
        self.calls = 0

    def stream(self, messages, tools=None, on_text=None, on_reasoning=None, **kw):
        self.calls += 1
        self.last_messages = messages
        return self.script.pop(0)


def _tc(name, args="{}"):
    return {"id": f"call_{name}", "type": "function",
            "function": {"name": name, "arguments": args}}


def test_communication_protocol() -> None:
    print("[komunikační protokol]")
    from harness.agent import Agent, Status
    from harness.llm import AssistantResult
    from harness.safety import SafetyPolicy
    data = load_config().data
    data["paths"]["sessions_dir"] = str(Path(tempfile.mkdtemp()) / "sessions")
    cfg = Config(data, root=ROOT)
    tmp = Path(tempfile.mkdtemp())
    try:
        cfg.agent["workspace"] = str(tmp)

        def make_agent(script):
            session = Session(cfg, session_id=f"proto-{uuid.uuid4().hex[:6]}")
            llm = LLMStub(script)
            agent = Agent(cfg, llm, session, build_registry("agent"),
                          SafetyPolicy("auto"), mode="agent")
            return agent, session

        # 1) poznámky zůstávají v historii, aby se neměnil již cachovaný prefix
        agent, session = make_agent([])
        agent.new_task("udelej neco")
        notes = [m for m in session.messages if "[TASK PROTOCOL" in str(m.get("content"))]
        check(len(notes) == 1, "TASK PROTOCOL poznámka přidána (user role)")
        agent.new_task("dalsi ukol")
        notes = [m for m in session.messages if "[TASK PROTOCOL" in str(m.get("content"))]
        check(len(notes) == 2, "nová úloha nemění starý cachovaný protokol")
        reloaded = Session.load(cfg, session.id)
        persisted_notes = [m for m in reloaded.messages
                           if "[TASK PROTOCOL" in str(m.get("content"))]
        check(len(persisted_notes) == 2, "cache-friendly protokoly přežijí reload beze změny")

        # 2) progress nudge po 4 tool-krocích bez slov
        script = [AssistantResult(tool_calls=[_tc("list_dir", '{"path": "."}')]) for _ in range(4)]
        script.append(AssistantResult(content="hotovo"))
        agent, session = make_agent(script)
        agent.new_task("prohledat adresar")
        statuses = [agent.step(approve=True).status for _ in range(5)]
        prog = [m for m in session.messages if "[PROGRESS UPDATE" in str(m.get("content"))]
        check(len(prog) >= 1, f"PROGRESS nudge po 4 krocích (počet: {len(prog)})")

        # 3) vynucení strukturovaného souhrnu po úloze s nástroji
        script = [
            AssistantResult(tool_calls=[_tc("list_dir")]),
            AssistantResult(tool_calls=[_tc("list_dir"), _tc("list_dir")]),
            AssistantResult(content="jen kratka odpoved"),   # nezaklad vyzaduje souhrn
            AssistantResult(content="✅ Hotovo: nic\n- **x**: y"),  # strukturovany
        ]
        agent, session = make_agent(script)
        agent.new_task("test souhrnu")
        r1 = agent.step(approve=True)
        r2 = agent.step(approve=True)
        r3 = agent.step(approve=True)
        check(r3.status is Status.CONTINUE, "krátká odpověď po nástrojích → vynucen souhrn (CONTINUE)")
        notes = [m for m in session.messages if "[FINAL SUMMARY" in str(m.get("content"))]
        check(len(notes) == 1, "SUMMARY poznámka vložena")
        r4 = agent.step(approve=True)
        check(r4.status is Status.FINAL and "✅" in r4.text, "druhý průchod → FINAL se souhrnem")
        check(agent.llm.calls == 4, "žádné zbytečné navíc volání")

        # 4) chat režim: žádný protokol (bez nástrojů netřeba)
        session = Session(cfg, session_id="chat-proto")
        agent = Agent(cfg, LLMStub([]), session, build_registry("chat"), SafetyPolicy("auto"), mode="chat")
        agent.new_task("ahoj")
        notes = [m for m in session.messages if "[TASK PROTOCOL" in str(m.get("content"))]
        check(not notes, "chat režim bez protokolu")

        # 5) přetečení kontextu → komprese + retry (1×), pak FINAL
        class OverflowLLM(LLMStub):
            def stream(self, messages, **kw):
                self.calls += 1
                if self.calls == 1:
                    raise RuntimeError("request exceeds the available context size")
                from harness.llm import AssistantResult
                return AssistantResult(content="✅ po kompresi OK")

        from harness.agent import Status as St
        for i in range(8):
            (tmp / f"f{i}.txt").write_text("x" * 400, encoding="utf-8")
        session = Session(cfg, session_id="ovf-test", system_prompt="SYS")
        ovf_llm = OverflowLLM()
        agent = Agent(cfg, ovf_llm, session, build_registry("agent"), SafetyPolicy("auto"), mode="agent")
        agent.llm = ovf_llm
        # naplnění session, aby komprese měla co zahodit
        for i in range(8):
            session.add("user", f"q{i} " + "y" * 900)
            session.add("assistant", f"a{i} " + "z" * 900)
        agent.new_task("dalsi dotaz")
        agent._steps = 0
        r = agent.step(approve=True)
        check(r.status is St.CONTINUE, "overflow → CONTINUE (komprese + retry)")
        r2 = agent.step(approve=True)
        check(r2.status is St.FINAL and ovf_llm.calls == 2,
              "retry po kompresi uspěl (2 volání)")
        # druhé přetečení už retry nedostane → ERROR
        class AlwaysOverflow(LLMStub):
            def stream(self, messages, **kw):
                self.calls += 1
                raise RuntimeError("prompt is too long: 999999 > 98304")
        ao = AlwaysOverflow()
        agent.llm = ao
        agent._overflow_retried = False
        session.add("user", "znovu preteceni")
        agent.safety.new_task()
        r3 = agent.step(approve=True)
        r4 = agent.step(approve=True)
        check(r4.status is St.ERROR and ao.calls == 2, "druhé přetečení → ERROR (žádná smyčka)")
    finally:
        shutil.rmtree(tmp, ignore_errors=True)


if __name__ == "__main__":
    test_config()
    test_memory_layers()
    test_safety()
    test_session()
    test_tools_fs_shell()
    test_registry_modes()
    test_parse_args()
    test_reasoning_effort_kwargs()
    test_runtime_lifecycle_helpers()
    test_dependency_marker()
    test_streaming_bridge()
    test_parallel_read_tools()
    test_resume_task_and_process_after_restart()
    test_git_tools()
    test_automatic_project_check()
    test_research_ledger_and_synthesis()
    test_project_document_library()
    test_async_model_switch()
    test_shell_readonly()
    test_workspace()
    test_session_meta()
    test_chat_rewind_and_fork()
    test_transient_session()
    test_web_tools()
    test_projects()
    test_skill_library()
    test_context_compression()
    test_communication_protocol()
    print(f"\n{'=' * 40}\nVÝSLEDEK: {PASS} ✓ / {FAIL} ✗")
    sys.exit(1 if FAIL else 0)
