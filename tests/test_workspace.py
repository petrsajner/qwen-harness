"""End-to-end service contracts without a GPU or browser process."""
from __future__ import annotations

import copy
import io
import json
from pathlib import Path
import tempfile
import threading
import time
import unittest
from types import SimpleNamespace

from fastapi.testclient import TestClient
from PIL import Image

from harness.application import ApplicationService
from harness.app_operations import perform_action
from harness.app_storage import export_project, import_project
from harness.changes import ChangeJournal
from harness.config import Config, load_config
from harness.documents import read_document_content
from harness.llm import AssistantResult, LLMClient
from harness.projects import Projects
from harness.session import Session
from harness.web_api import create_app


def wait_for(predicate, timeout=5):
    deadline = time.monotonic() + timeout
    while time.monotonic() < deadline:
        if predicate():
            return
        time.sleep(0.01)
    raise AssertionError("Timed out waiting for operation")


class Model:
    calls = []
    gate = threading.Event()

    def __init__(self, cfg):
        self.cfg = cfg

    def stream(self, messages, **kwargs):
        self.calls.append((copy.deepcopy(messages), copy.deepcopy(self.cfg.data)))
        should_stop = kwargs.get("should_stop", lambda: False)
        on_text = kwargs.get("on_text", lambda text: None)
        content = [m.get("content", "") for m in messages if m.get("role") == "user"]
        content = " ".join(str(c) for c in content)
        if "hold-task" in content and "clarification" not in content:
            on_text("Partial answer.")
            while not self.gate.wait(0.01):
                if should_stop():
                    return AssistantResult(content="Partial answer.", stopped=True)
        on_text("Finished answer.")
        return AssistantResult(content="Finished answer.", usage={"prompt_tokens": 42, "completion_tokens": 8})


class WorkspaceTests(unittest.TestCase):
    def setUp(self):
        self.temp = tempfile.TemporaryDirectory()
        self.root = Path(self.temp.name)
        data = copy.deepcopy(load_config().data)
        data["agent"]["workspace"] = None
        data["agent"]["autonomy"] = "auto"
        data["work_mode"] = "discussion"
        data["hardware"]["vram_gb"] = 32
        self.cfg = Config(data, self.root)
        Model.calls = []
        Model.gate = threading.Event()
        self.service = ApplicationService(self.cfg, llm_factory=Model, manage_model=False)
        self.client = TestClient(create_app(self.cfg, service=self.service))
        self.client.__enter__()

    def tearDown(self):
        Model.gate.set()
        self.client.__exit__(None, None, None)
        self.temp.cleanup()

    def new_chat(self, **kwargs):
        return self.client.post("/api/sessions", json=kwargs).json()["session_id"]

    def completed(self, key):
        wait_for(lambda: self.service.store.job(key)["status"] in ("complete", "stopped", "failed"))

    def test_fork_keeps_document_after_original_chat_deleted(self):
        sid = self.new_chat()
        session = self.service.session(sid)
        session.persist()
        source = session.dir / "attachments" / "notes.txt"
        source.parent.mkdir(parents=True, exist_ok=True)
        source.write_text("Keep this document", encoding="utf-8")
        item = self.service.store.register_file(source, sid, "attachment", "notes.txt")
        session.add("user", "Read " + str(source))
        session.messages[-1]["attachments"] = [item["id"]]
        session._rewrite_jsonl()
        result = perform_action(self.service, sid, "fork", {})
        fork = self.service.session(result["session_id"])
        copied = self.service.store.file(fork.messages[-1]["attachments"][0])
        perform_action(self.service, sid, "delete", {})
        self.assertEqual(Path(copied["path"]).read_text(encoding="utf-8"), "Keep this document")
        self.assertIn(copied["path"], fork.messages[-1]["content"])
        self.assertNotEqual(copied["id"], item["id"])

    def test_navigation_does_not_rebind_running_task_and_submit_is_idempotent(self):
        first, second = self.new_chat(), self.new_chat()
        job = {"text": "hold-task", "request_id": "one"}
        self.client.post(f"/api/sessions/{first}/submit", json=job).raise_for_status()
        wait_for(lambda: bool(self.service.active))
        self.client.post(f"/api/sessions/{first}/submit", json=job).raise_for_status()
        self.client.post(f"/api/sessions/{second}/select").raise_for_status()
        self.assertEqual(self.service.active["session_id"], first)
        Model.gate.set()
        self.completed("one")
        self.assertEqual(len([m for m in self.service.session(first).messages if m.get("content") == "hold-task"]), 1)
        self.assertFalse(self.service.session(second).messages)

    def test_steering_and_queue_capture_settings(self):
        sid = self.new_chat()
        self.client.post(f"/api/sessions/{sid}/submit", json={"text": "hold-task", "request_id": "first"})
        wait_for(lambda: len(Model.calls) == 1)
        self.client.patch("/api/settings", json={"thinking": "low"})
        self.client.post(f"/api/sessions/{sid}/submit", json={"text": "clarification", "request_id": "steer"})
        self.completed("first")
        self.assertEqual(self.service.store.job("steer")["status"], "complete")
        self.assertEqual(Model.calls[-1][1]["reasoning_effort"], "xhigh")
        self.client.post(f"/api/sessions/{sid}/submit", json={"text": "next", "request_id": "next", "delivery": "queue"})
        self.completed("next")
        self.assertEqual(Model.calls[-1][1]["reasoning_effort"], "low")
        self.assertTrue(any(m.get("content") == "Partial answer." for m in self.service.session(sid).messages))

    def test_attachment_payload_names_preview_and_reload(self):
        sid = self.new_chat()
        buffer = io.BytesIO()
        Image.new("RGB", (20, 16), "green").save(buffer, format="PNG")
        uploaded = self.client.post(f"/api/sessions/{sid}/attachments", files={"file": ("náhled.png", buffer.getvalue(), "image/png")}).json()
        self.client.post(f"/api/sessions/{sid}/actions/draft", json={"text": "draft", "attachments": [uploaded]}).raise_for_status()
        self.assertEqual(self.client.get(f"/api/sessions/{sid}").json()["draft"]["attachments"][0]["name"], "náhled.png")
        self.client.post(f"/api/sessions/{sid}/submit", json={"text": "", "attachments": [uploaded["id"]], "request_id": "image"}).raise_for_status()
        self.completed("image")
        user = next(m for m in Model.calls[0][0] if isinstance(m.get("content"), list))
        self.assertTrue(any(p.get("type") == "image_url" and p["image_url"]["url"].startswith("data:image/png;base64,") for p in user["content"]))
        response = self.client.get(f"/api/sessions/{sid}").json()
        files = [f for m in response["messages"] if m["role"] == "user" for f in m["files"]]
        self.assertEqual(len(files), 1)
        self.assertEqual(files[0]["name"], "náhled.png")
        self.assertEqual(self.client.get(files[0]["url"]).content, buffer.getvalue())
        loaded = Session.load(self.cfg, sid)
        self.assertTrue(Path(next(m for m in loaded.messages if m.get("images"))["images"][0]).is_file())

    def test_stop_is_visible_and_does_not_release_queued_work(self):
        sid = self.new_chat()
        self.client.post(f"/api/sessions/{sid}/submit", json={"text": "hold-task", "request_id": "hold"})
        wait_for(lambda: bool(Model.calls))
        self.client.post(f"/api/sessions/{sid}/submit", json={"text": "next", "request_id": "queued", "delivery": "queue"})
        self.client.post(f"/api/sessions/{sid}/actions/stop", json={})
        self.completed("hold")
        self.assertEqual(self.service.store.job("hold")["status"], "stopped")
        self.assertEqual(self.service.store.job("queued")["status"], "queued")
        self.assertTrue(self.service.queue_paused)

    def test_project_move_reconfigures_chat_and_keeps_other_chat_unchanged(self):
        source = self.client.post("/api/projects", json={"name": "Project A", "mode": "writing"}).json()
        other = self.new_chat()
        sid = source["session_id"]
        result = self.client.post(f"/api/sessions/{sid}/actions/move", json={"project_id": None})
        result.raise_for_status()
        self.assertIsNone(self.service.session(sid).meta["workspace"])
        self.assertEqual(self.service.session(other).meta["work_mode"], "discussion")
        self.assertIsNone(self.service.config_for(self.service.session(sid)).agent["workspace"])

    def test_events_replay_and_no_internal_ids_in_model_payload(self):
        sid = self.new_chat()
        before = self.service.store.sequence()
        self.client.post(f"/api/sessions/{sid}/submit", json={"text": "hello", "request_id": "events"})
        self.completed("events")
        events = self.service.store.after(before, 1000)
        self.assertTrue(any(e["kind"] == "message" for e in events))
        self.assertEqual([e["seq"] for e in events], sorted(e["seq"] for e in events))
        self.assertFalse(any("id" in m or "run_id" in m for m in Model.calls[0][0]))

    def test_document_ranges_and_word_structure(self):
        import openpyxl
        from docx import Document
        wb = openpyxl.Workbook()
        ws = wb.active
        ws["A150"] = "late row"
        ws["B150"] = "=1+2"
        target = self.root / "book.xlsx"
        wb.save(target)
        text = read_document_content(target, cell_range="A150:B150", formulas=True)
        self.assertIn("late row", text)
        self.assertIn("=1+2", text)
        doc = Document()
        doc.add_paragraph("before")
        doc.add_table(rows=1, cols=1).cell(0, 0).text = "table"
        doc.add_paragraph("after")
        word = self.root / "word.docx"
        doc.save(word)
        text = read_document_content(word)
        self.assertLess(text.index("before"), text.index("table"))
        self.assertLess(text.index("table"), text.index("after"))

    def test_checkpoint_conflict_and_portable_project(self):
        project = Projects(self.cfg).create_new("Portable")
        workspace = Path(project["path"])
        file = workspace / "notes.txt"
        file.write_text("original")
        sid = self.service.new_session(str(workspace), "writing").id
        session = self.service.session(sid)
        session.add("user", "Remember this project")
        journal = ChangeJournal(session, workspace)
        cp = journal.create_checkpoint("before edit")
        file.write_text("model edit")
        journal.reconcile_workspace()
        file.write_text("later user edit")
        result = journal.undo(cp)
        self.assertTrue(result["errors"])
        self.assertEqual(file.read_text(), "later user edit")
        archive = export_project(self.cfg, project, self.root / "portable.zip")
        imported = import_project(self.cfg, archive)
        self.assertNotEqual(imported["path"], project["path"])
        self.assertEqual((Path(imported["path"]) / "notes.txt").read_text(), "later user edit")
        sessions = Session.list_sessions(self.cfg)
        self.assertTrue(any(s["workspace"] == imported["path"] for s in sessions))

    def test_portable_document_attachments_and_decision_links(self):
        project = Projects(self.cfg).create_new("Document project")
        session = self.service.new_session(project["path"], "discussion")
        target = session.dir / "attachments" / "document.txt"
        target.parent.mkdir(parents=True, exist_ok=True)
        target.write_text("portable document")
        file = self.service.store.register_file(target, session.id, "attachment", name="document.txt")
        message = session.add("user", f"Read this document: {target}")
        message["attachments"] = [file["id"]]
        session._rewrite_jsonl()
        from harness.decisions import DecisionStore
        DecisionStore(project["path"]).save("Keep the chosen format", session.id, "accepted")
        archive = export_project(self.cfg, project, self.root / "documents.zip")
        imported = import_project(self.cfg, archive)
        imported_session = next(s for s in Session.list_sessions(self.cfg) if s["workspace"] == imported["path"])
        loaded = Session.load(self.cfg, imported_session["id"])
        copied = self.service.store.file(loaded.messages[0]["attachments"][0])
        self.assertTrue(Path(copied["path"]).is_file())
        self.assertNotEqual(copied["path"], str(target))
        self.assertIn(copied["path"], loaded.messages[0]["content"])
        self.assertEqual(DecisionStore(imported["path"]).list()[0]["source_session"], loaded.id)

    def test_recovery_keeps_partial_text_and_marks_unknown_tool_outcome(self):
        sid = self.new_chat()
        session = self.service.session(sid)
        cfg = self.service.config_for(session)
        session.add("user", "recover-task")
        session.add("assistant", "", tool_calls=[{"id": "crash-call", "type": "function",
                    "function": {"name": "write_file", "arguments": '{"path":"already.txt","content":"done"}'}}])
        (session.dir / "already.txt").write_text("done")
        partial = {"text": "Visible partial text", "reasoning": "", "run_id": "recover"}
        (session.dir / "run-live.json").write_text(json.dumps(partial))
        job = {"id": "recover", "session_id": sid, "text": "recover-task", "attachments": [],
               "config": cfg.data, "settings": self.service.preferences, "created": time.time()}
        self.service.store.save_job(job, "running")
        self.service.close()
        recovered = ApplicationService(self.cfg, llm_factory=Model, manage_model=False)
        try:
            self.assertEqual(recovered.store.job("recover")["status"], "interrupted")
            recovered.resume(sid)
            wait_for(lambda: recovered.store.job("recover")["status"] == "complete")
            history = recovered.session(sid).messages
            self.assertTrue(any(m.get("content") == partial["text"] for m in history))
            self.assertEqual(len([m for m in history if m.get("tool_call_id") == "crash-call"]), 1)
            self.assertEqual((session.dir / "already.txt").read_text(), "done")
        finally:
            recovered.close()

    def test_compression_reads_middle_and_uses_discussion_focus(self):
        from harness.context import summarize_messages
        calls = []
        class Summarizer:
            cfg = self.cfg
            def stream(self, messages, **kwargs):
                calls.append(messages[-1]["content"])
                return AssistantResult(content="Decisions and original references preserved.")
        text = "A" * 190000 + "MIDDLE_REQUIREMENT_9281" + "Z" * 190000
        summarize_messages(Summarizer(), [{"id": "source-message", "role": "user", "content": text}])
        self.assertTrue(any("MIDDLE_REQUIREMENT_9281" in prompt for prompt in calls))
        self.assertIn("discussion", calls[0])
        self.assertNotIn("coding agent", calls[0])

    def test_pdf_vision_page_and_word_edit_preserves_style(self):
        from reportlab.pdfgen.canvas import Canvas
        from docx import Document
        from harness.tools.base import AgentContext
        from harness.tools.document_edit import EditWordTool, ViewDocumentPageTool
        session = self.service.new_session(work_mode="writing")
        ctx = AgentContext(self.cfg, session, workspace=self.root)
        ctx.changes = ChangeJournal(session, self.root)
        path = self.root / "scan.pdf"
        canvas = Canvas(str(path))
        canvas.setFillColorRGB(0, 1, 0)
        canvas.rect(10, 10, 200, 200, fill=True)
        canvas.save()
        ViewDocumentPageTool().run(ctx, str(path))
        self.assertTrue(ctx.pending_images[0].is_file())
        word = self.root / "edit.docx"
        doc = Document()
        paragraph = doc.add_paragraph()
        paragraph.add_run("Keep ").bold = True
        paragraph.add_run("old wording").italic = True
        doc.add_table(rows=1, cols=1).cell(0, 0).text = "table retained"
        doc.save(word)
        EditWordTool().run(ctx, str(word), "old wording", "new wording")
        reloaded = Document(word)
        self.assertTrue(reloaded.paragraphs[0].runs[0].bold)
        self.assertTrue(reloaded.paragraphs[0].runs[1].italic)
        self.assertEqual(reloaded.tables[0].cell(0, 0).text, "table retained")
        doc = Document()
        doc.add_paragraph("old old")
        doc.save(word)
        EditWordTool().run(ctx, str(word), "old", "old-new", replace_all=True)
        self.assertEqual(Document(word).paragraphs[0].text, "old-new old-new")


class TransportTests(unittest.TestCase):
    def test_stop_before_any_bytes_returns_quickly(self):
        cfg = load_config()
        client = LLMClient.__new__(LLMClient)
        client.cfg, client.model_name = cfg, "test"
        blocked = threading.Event()
        closed = threading.Event()
        class Stream:
            def __iter__(self):
                blocked.wait(3)
                return iter([])
            def close(self):
                closed.set()
                blocked.set()
        client.client = SimpleNamespace(chat=SimpleNamespace(completions=SimpleNamespace(create=lambda **kw: Stream())))
        stop = threading.Event()
        timer = threading.Timer(0.05, stop.set)
        timer.start()
        started = time.monotonic()
        result = client.stream([], should_stop=stop.is_set)
        self.assertTrue(result.stopped)
        self.assertLess(time.monotonic() - started, 0.8)
        self.assertTrue(closed.is_set())
        timer.join()


if __name__ == "__main__":
    unittest.main()
