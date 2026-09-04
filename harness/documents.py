"""Export strukturovaného textu do Markdown, DOCX a PDF."""
from __future__ import annotations

import re
from html import escape
from pathlib import Path


def _safe_filename(name: str, suffix: str) -> str:
    stem = re.sub(r'[<>:"/\\|?*]+', "-", Path(name or "dokument").stem).strip(" .-")
    return (stem or "dokument") + suffix


def export_document(content: str, output_dir: Path, filename: str,
                    fmt: str, title: str = "") -> Path:
    fmt = fmt.lower()
    output_dir.mkdir(parents=True, exist_ok=True)
    target = document_target(output_dir, filename, fmt)
    if fmt == "markdown":
        target.write_text(content, encoding="utf-8", newline="\n")
    elif fmt == "docx":
        _write_docx(target, content, title)
    else:
        _write_pdf(target, content, title)
    return target


def document_target(output_dir: Path, filename: str, fmt: str) -> Path:
    suffix = {"markdown": ".md", "docx": ".docx", "pdf": ".pdf"}.get(fmt.lower())
    if suffix is None:
        raise ValueError("format must be markdown, docx, or pdf")
    return output_dir / _safe_filename(filename, suffix)


def _blocks(content: str):
    paragraph: list[str] = []
    for raw in content.splitlines():
        line = raw.rstrip()
        if not line:
            if paragraph:
                yield "paragraph", " ".join(paragraph)
                paragraph = []
            continue
        heading = re.match(r"^(#{1,4})\s+(.+)$", line)
        bullet = re.match(r"^\s*[-*]\s+(.+)$", line)
        numbered = re.match(r"^\s*\d+[.)]\s+(.+)$", line)
        table = re.match(r"^\s*\|(.+)\|\s*$", line)
        if heading or bullet or numbered or table or line.strip() == "---":
            if paragraph:
                yield "paragraph", " ".join(paragraph)
                paragraph = []
            if line.strip() == "---":
                continue
            if heading:
                yield f"heading{len(heading.group(1))}", heading.group(2)
            elif bullet:
                yield "bullet", bullet.group(1)
            elif numbered:
                yield "number", numbered.group(1)
            else:
                cells = [cell.strip() for cell in table.group(1).split("|")]
                if not all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells):
                    yield "table_row", cells
        else:
            paragraph.append(line)
    if paragraph:
        yield "paragraph", " ".join(paragraph)


def _write_docx(path: Path, content: str, title: str) -> None:
    from docx import Document
    document = Document()
    if title:
        document.add_heading(title, level=0)
    for kind, text in _blocks(content):
        if kind.startswith("heading"):
            document.add_heading(text, level=min(int(kind[-1]), 4))
        elif kind == "bullet":
            document.add_paragraph(text, style="List Bullet")
        elif kind == "number":
            document.add_paragraph(text, style="List Number")
        elif kind == "table_row":
            document.add_paragraph(" | ".join(text))
        else:
            document.add_paragraph(text)
    document.save(path)


def _write_pdf(path: Path, content: str, title: str) -> None:
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_LEFT
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import mm
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

    font = "Helvetica"
    family = {
        "normal": Path(r"C:\Windows\Fonts\arial.ttf"),
        "bold": Path(r"C:\Windows\Fonts\arialbd.ttf"),
        "italic": Path(r"C:\Windows\Fonts\ariali.ttf"),
        "boldItalic": Path(r"C:\Windows\Fonts\arialbi.ttf"),
    }
    if all(source.is_file() for source in family.values()):
        names = {
            "normal": "QwenUnicode",
            "bold": "QwenUnicode-Bold",
            "italic": "QwenUnicode-Italic",
            "boldItalic": "QwenUnicode-BoldItalic",
        }
        for variant, source in family.items():
            pdfmetrics.registerFont(TTFont(names[variant], str(source)))
        pdfmetrics.registerFontFamily(
            "QwenUnicode", normal=names["normal"], bold=names["bold"],
            italic=names["italic"], boldItalic=names["boldItalic"])
        font = "QwenUnicode"
    styles = getSampleStyleSheet()
    body = ParagraphStyle("BodyUnicode", parent=styles["BodyText"], fontName=font,
                          fontSize=10.5, leading=14, alignment=TA_LEFT, spaceAfter=5)
    headings = {
        level: ParagraphStyle(f"H{level}Unicode", parent=styles[f"Heading{min(level, 4)}"],
                              fontName=font, spaceBefore=8, spaceAfter=5)
        for level in range(1, 5)
    }
    table_body = ParagraphStyle("TableUnicode", parent=body, fontSize=8.5,
                                leading=11, spaceAfter=0)
    document = SimpleDocTemplate(str(path), pagesize=A4, rightMargin=18 * mm,
                                 leftMargin=18 * mm, topMargin=18 * mm,
                                 bottomMargin=18 * mm)
    blocks = list(_blocks(content))
    story = []
    if title:
        story.extend([Paragraph(_inline_markdown(title), headings[1]), Spacer(1, 3 * mm)])
    index = 0
    while index < len(blocks):
        kind, text = blocks[index]
        if kind == "table_row":
            rows = []
            while index < len(blocks) and blocks[index][0] == "table_row":
                rows.append(blocks[index][1])
                index += 1
            columns = max(len(row) for row in rows)
            normalized = [row + [""] * (columns - len(row)) for row in rows]
            data = [[Paragraph(_inline_markdown(cell), table_body) for cell in row]
                    for row in normalized]
            table = Table(data, colWidths=[document.width / columns] * columns,
                          repeatRows=1, hAlign="LEFT")
            table.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#E8EEF5")),
                ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#9AA6B2")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 4),
                ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                ("TOPPADDING", (0, 0), (-1, -1), 3),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
            ]))
            story.extend([table, Spacer(1, 2 * mm)])
            continue
        marked = _inline_markdown(text)
        if kind.startswith("heading"):
            story.append(Paragraph(marked, headings[int(kind[-1])]))
        elif kind == "bullet":
            story.append(Paragraph(marked, body, bulletText="•"))
        elif kind == "number":
            story.append(Paragraph(marked, body))
        else:
            story.append(Paragraph(marked, body))
        index += 1
    document.build(story)


def _escape(text: str) -> str:
    return escape(text, quote=True)


def _inline_markdown(text: str) -> str:
    replacements = {
        "⁺": "+", "⁻": "-", "‑": "-", "✅": "", "🔍": "",
        "📋": "", "⏸️": "", "⏸": "", "\ufe0f": "",
    }
    normalized = text
    for source, target in replacements.items():
        normalized = normalized.replace(source, target)
    marked = _escape(normalized)
    marked = re.sub(r"\[([^\]]+)\]\((https?://[^\s)]+)\)",
                    r'<a href="\2" color="#165D9C">\1</a>', marked)
    marked = re.sub(r"`([^`]+)`", r'<font color="#394B59">\1</font>', marked)
    marked = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", marked)
    marked = re.sub(r"(?<!\*)\*([^*]+)\*(?!\*)", r"<i>\1</i>", marked)
    return marked


def _format_markdown_table(rows: list[list[str]]) -> str:
    if not rows:
        return ""
    max_cols = max(len(r) for r in rows)
    if max_cols == 0:
        return ""
    padded = [r + [""] * (max_cols - len(r)) for r in rows]
    header = padded[0]
    out = ["| " + " | ".join(header) + " |"]
    out.append("| " + " | ".join(["---"] * max_cols) + " |")
    for r in padded[1:]:
        out.append("| " + " | ".join(r) + " |")
    return "\n".join(out)


def read_document_content(path: Path, max_chars: int = 40_000, sheet: str | None = None) -> str:
    """Extrahuje čistý text a tabulky z dokumentů (.docx, .pdf, .xlsx, .csv)."""
    p = Path(path).resolve()
    if not p.is_file():
        raise FileNotFoundError(f"Dokument nebyl nalezen: {p}")

    suffix = p.suffix.lower()
    if suffix == ".docx":
        import docx
        doc = docx.Document(p)
        parts = [f"=== Word Document: {p.name} ==="]
        for para in doc.paragraphs:
            txt = para.text.strip()
            if not txt:
                continue
            if para.style and para.style.name and para.style.name.startswith("Heading"):
                lvl = para.style.name.replace("Heading", "").strip() or "2"
                h_num = int(lvl) if lvl.isdigit() else 2
                parts.append(f"{'#' * min(4, h_num)} {txt}")
            else:
                parts.append(txt)
        for t in doc.tables:
            rows = []
            for row in t.rows:
                rows.append([cell.text.strip().replace("\n", " ") for cell in row.cells])
            if rows:
                parts.append(_format_markdown_table(rows))
        res = "\n\n".join(parts)

    elif suffix == ".pdf":
        import pypdf
        reader = pypdf.PdfReader(p)
        parts = [f"=== PDF Document: {p.name} ({len(reader.pages)} stran) ==="]
        for idx, page in enumerate(reader.pages, 1):
            txt = page.extract_text() or ""
            if txt.strip():
                parts.append(f"--- Strana {idx} ---\n{txt.strip()}")
        res = "\n\n".join(parts)

    elif suffix in (".xlsx", ".xlsm"):
        import openpyxl
        wb = openpyxl.load_workbook(p, data_only=True)
        parts = [f"=== Excel sešit: {p.name} (listy: {', '.join(wb.sheetnames)}) ==="]
        sheets = [sheet] if sheet and sheet in wb.sheetnames else wb.sheetnames[:5]
        for s_name in sheets:
            ws = wb[s_name]
            rows = []
            for r in ws.iter_rows(values_only=True):
                if any(c is not None for c in r):
                    rows.append([str(c) if c is not None else "" for c in r])
            if rows:
                parts.append(f"### List: `{s_name}`\n" + _format_markdown_table(rows[:100]))
                if len(rows) > 100:
                    parts.append(f"*... [vynecháno dalších {len(rows) - 100} řádků]*")
            else:
                parts.append(f"### List: `{s_name}` (prázdný)")
        res = "\n\n".join(parts)

    elif suffix == ".csv":
        import csv
        with open(p, "r", encoding="utf-8", errors="replace") as f:
            rows = list(csv.reader(f))
            parts = [f"=== CSV: {p.name} ({len(rows)} řádků) ==="]
            if rows:
                parts.append(_format_markdown_table(rows[:100]))
            res = "\n\n".join(parts)
    else:
        res = p.read_text(encoding="utf-8", errors="replace")

    if len(res) > max_chars:
        half = max_chars // 2
        omitted = len(res) - max_chars
        res = res[:half] + f"\n\n... [dokument zkrácen: ~{omitted} znaků vynecháno] ...\n\n" + res[-half:]
    return res


def edit_spreadsheet_content(path: Path, action: str, sheet: str | None = None,
                             data: Any = None, title: str | None = None) -> str:
    """Vytváří a edituje Excel tabulky (.xlsx)."""
    import openpyxl
    p = Path(path).resolve()
    p.parent.mkdir(parents=True, exist_ok=True)
    action = action.lower().strip()

    if action == "create":
        wb = openpyxl.Workbook()
        ws = wb.active
        if title:
            ws.title = title[:31]
        if isinstance(data, list):
            for row in data:
                if isinstance(row, (list, tuple)):
                    ws.append(list(row))
        wb.save(p)
        return f"OK: Vytvořen nový Excel sešit `{p.name}` s listem `{ws.title}`."

    if not p.is_file():
        raise FileNotFoundError(f"Excel soubor nebyl nalezen: {p}. Pro vytvoření nového zadejte action='create'.")

    wb = openpyxl.load_workbook(p)
    ws = wb[sheet] if sheet and sheet in wb.sheetnames else wb.active

    if action == "list_sheets":
        return f"Listy v sešitu `{p.name}`: " + ", ".join(f"`{s}`" for s in wb.sheetnames)

    if action == "create_sheet":
        s_title = (title or sheet or "NovyList")[:31]
        new_ws = wb.create_sheet(title=s_title)
        wb.save(p)
        return f"OK: Vytvořen nový list `{new_ws.title}` v sešitu `{p.name}`."

    if action == "append_rows":
        if not isinstance(data, list):
            raise ValueError("Pro action='append_rows' musí být data seznam řádků (list of lists).")
        for row in data:
            if isinstance(row, (list, tuple)):
                ws.append(list(row))
        wb.save(p)
        return f"OK: Přidáno {len(data)} řádků do listu `{ws.title}`."

    if action == "update_cells":
        if not isinstance(data, dict):
            raise ValueError("Pro action='update_cells' musí být data slovník buněk např. {'A1': 100, 'B1': '=SUM(A1:A5)'}.")
        for cell_coord, val in data.items():
            ws[cell_coord] = val
        wb.save(p)
        return f"OK: Aktualizováno {len(data)} buněk v listu `{ws.title}`."

    raise ValueError(f"Neznámá akce: {action}. Dostupné akce: create, list_sheets, create_sheet, append_rows, update_cells")

